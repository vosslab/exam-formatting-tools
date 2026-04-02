#!/usr/bin/env python3
"""Convert a DOCX exam file to the YAML format expected by odt_exam_builder.py.

Parses question text, multiple choice options, embedded images, and tables
from a Word DOCX file and writes structured YAML output.
"""

# Standard Library
import os
import re
import argparse

# Pip Modules
import yaml
import docx
import docx.oxml.ns


#============================================
def parse_args() -> argparse.Namespace:
	"""Parse command-line arguments.

	Returns:
		Parsed argument namespace.
	"""
	parser = argparse.ArgumentParser(
		description="Convert a DOCX exam to YAML for odt_exam_builder"
	)
	parser.add_argument(
		'-i', '--input', dest='input_file', required=True,
		help="Input DOCX file"
	)
	parser.add_argument(
		'-o', '--output', dest='output_file', required=True,
		help="Output YAML file"
	)
	parser.add_argument(
		'-d', '--image-dir', dest='image_dir', default='images',
		help="Directory to extract images into"
	)
	args = parser.parse_args()
	return args


#============================================
def extract_images(doc: docx.Document, image_dir: str) -> dict:
	"""Extract all images from the DOCX and save to image_dir.

	Maps relationship IDs to saved file paths so paragraphs with
	images can reference them.

	Args:
		doc: The loaded DOCX document.
		image_dir: Directory to save extracted images.

	Returns:
		Dict mapping paragraph index to image file path.
	"""
	os.makedirs(image_dir, exist_ok=True)
	# map relationship IDs to image parts
	rel_to_image = {}
	for rel_id, rel in doc.part.rels.items():
		if "image" in rel.reltype:
			rel_to_image[rel_id] = rel.target_part

	# find which paragraphs contain images and extract them
	para_images = {}
	image_count = 0
	for para_idx, para in enumerate(doc.paragraphs):
		# look for drawing elements containing image references
		drawings = para._element.findall('.//' + docx.oxml.ns.qn('w:drawing'))
		if not drawings:
			continue
		for drawing in drawings:
			# find the blip element which has the image relationship ID
			blips = drawing.findall('.//' + docx.oxml.ns.qn('a:blip'))
			for blip in blips:
				embed_id = blip.get(docx.oxml.ns.qn('r:embed'))
				if embed_id is None:
					continue
				if embed_id not in rel_to_image:
					continue
				image_part = rel_to_image[embed_id]
				# determine file extension from content type
				content_type = image_part.content_type
				ext = '.png'
				if 'jpeg' in content_type or 'jpg' in content_type:
					ext = '.jpg'
				elif 'gif' in content_type:
					ext = '.gif'
				elif 'emf' in content_type:
					ext = '.emf'
				elif 'wmf' in content_type:
					ext = '.wmf'
				image_count += 1
				image_filename = f"exam_image_{image_count:02d}{ext}"
				image_path = os.path.join(image_dir, image_filename)
				# write image bytes to file
				with open(image_path, 'wb') as f:
					f.write(image_part.blob)
				para_images[para_idx] = image_path
	return para_images


#============================================
def extract_table_data(table: docx.table.Table) -> dict:
	"""Extract table data into a dict with columns and rows.

	Args:
		table: A python-docx Table object.

	Returns:
		Dict with 'columns' (list of header strings) and 'rows'
		(list of lists of cell strings).
	"""
	rows = []
	for row in table.rows:
		cells = [cell.text.strip() for cell in row.cells]
		rows.append(cells)
	# first row is header
	columns = rows[0]
	data_rows = rows[1:]
	result = {
		'columns': columns,
		'rows': data_rows,
	}
	return result


#============================================
def parse_inline_choices(text: str) -> tuple:
	"""Parse choices from inline text with A)/A./A: patterns.

	Handles choices separated by newlines or tabs.

	Args:
		text: The full paragraph text that may contain inline choices.

	Returns:
		Tuple of (question_text, list_of_choice_strings) or
		(original_text, None) if no choices found.
		Choice strings are plain text without letter prefixes.
	"""
	# pattern: letter followed by ) or . or : then text
	# choices typically start after a newline
	choice_pattern = re.compile(
		r'\n\s*([A-E][).\s])\s*(.+?)(?=\n\s*[A-E][).\s]|\Z)',
		re.DOTALL
	)
	matches = list(choice_pattern.finditer(text))
	if len(matches) < 2:
		return (text, None)
	# question text is everything before first choice
	first_match_start = matches[0].start()
	question_text = text[:first_match_start].strip()
	# extract choices without letter prefixes
	choices = []
	for match in matches:
		choice_text = match.group(2).strip()
		choices.append(choice_text)
	result = (question_text, choices)
	return result


#============================================
def detect_question_number(text: str) -> int:
	"""Extract the question number from the beginning of text.

	Args:
		text: Paragraph text that may start with a number.

	Returns:
		The question number, or 0 if not found.
	"""
	match = re.match(r'^\s*(\d+)\s*[).]', text)
	if match is None:
		return 0
	num = int(match.group(1))
	return num


#============================================
def parse_docx_questions(doc: docx.Document, para_images: dict) -> list:
	"""Parse all questions from the DOCX paragraphs.

	Handles multiple patterns:
	- Inline choices (choices embedded in paragraph text with newlines)
	- List Paragraph choices (separate paragraphs with List Paragraph style)
	- Questions with images
	- Questions with tables

	Args:
		doc: The loaded DOCX document.
		para_images: Dict mapping paragraph index to image file path.

	Returns:
		List of question dicts in odt_exam_builder YAML format.
		Uses 'statement' key for question text, 'choices' as plain list.
	"""
	paragraphs = doc.paragraphs
	questions = []
	current_question = None
	# track which paragraph indices belong to tables
	# (python-docx intermixes tables and paragraphs in the body)
	# we need to find tables by checking the document body XML
	table_positions = {}
	body = doc.element.body
	element_idx = 0
	para_counter = 0
	table_counter = 0
	for child in body:
		tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
		if tag == 'p':
			para_counter += 1
		elif tag == 'tbl':
			# this table appears after paragraph (para_counter - 1)
			table_positions[para_counter] = table_counter
			table_counter += 1
		element_idx += 1

	i = 0
	while i < len(paragraphs):
		para = paragraphs[i]
		text = para.text.strip()
		style = para.style.name

		# skip empty paragraphs
		if not text and i not in para_images:
			# check if a table follows this empty paragraph
			# table_positions maps the para index after which the table appears
			i += 1
			continue

		# check if this paragraph starts a new question
		q_num = detect_question_number(text)
		if q_num > 0:
			# save previous question
			if current_question is not None:
				questions.append(current_question)
			# parse inline choices
			question_text, inline_choices = parse_inline_choices(text)
			current_question = {
				'statement': question_text,
			}
			if inline_choices is not None:
				current_question['choices'] = inline_choices
			else:
				# check for List Paragraph choices following
				list_choices = []
				continuation_text = []
				j = i + 1
				blank_count = 0
				while j < len(paragraphs):
					next_para = paragraphs[j]
					next_text = next_para.text.strip()
					next_style = next_para.style.name
					# check for image paragraphs
					if j in para_images:
						# image paragraph between question and choices
						current_question['image'] = para_images[j]
						j += 1
						blank_count = 0
						continue
					if next_style == 'List Paragraph' and next_text:
						list_choices.append(next_text)
						j += 1
						blank_count = 0
						continue
					if not next_text:
						blank_count += 1
						j += 1
						# stop at 3 consecutive blanks
						if blank_count >= 3:
							break
						continue
					# check if this is a new question
					if detect_question_number(next_text) > 0:
						break
					# otherwise it is continuation text
					blank_count = 0
					continuation_text.append(next_text)
					j += 1
				# attach continuation text to the question statement
				if continuation_text:
					current_question['statement'] += '\n' + '\n'.join(continuation_text)
				if list_choices:
					# strip letter prefixes from choices
					cleaned_choices = []
					for choice in list_choices:
						# check if has letter prefix (A) or A. or A)
						cleaned = re.sub(r'^[A-E][).\s]\s*', '', choice).strip()
						cleaned_choices.append(cleaned)
					current_question['choices'] = cleaned_choices
					i = j
					continue
			# check for image on the same paragraph or nearby
			if i in para_images and 'image' not in current_question:
				current_question['image'] = para_images[i]
			# check paragraph right after for image
			if (i + 1) in para_images and 'image' not in current_question:
				current_question['image'] = para_images[i + 1]
			i += 1
			continue

		# non-question paragraph that might be continuation text
		if current_question is not None:
			# check if it is continuation text for the current question
			if text and style != 'List Paragraph':
				# append to the question statement
				current_question['statement'] += '\n' + text
			# check for image
			if i in para_images:
				current_question['image'] = para_images[i]

		i += 1

	# save last question
	if current_question is not None:
		questions.append(current_question)

	# attach table data to questions that reference tables
	# find questions near table positions
	tables = doc.tables
	for para_idx, table_idx in table_positions.items():
		if table_idx >= len(tables):
			continue
		table_data = extract_table_data(tables[table_idx])
		# find the question closest to this table position
		# assign table to the most recently parsed question
		if questions:
			best_q = questions[-1]
			if 'table' not in best_q:
				best_q['table'] = table_data

	# post-process: parse choices from statement field if question has no choices
	for q in questions:
		if 'choices' in q:
			continue
		# statement is required, try parsing choices from it
		text = q['statement']
		q_text, choices = parse_inline_choices(text)
		if choices is not None:
			q['statement'] = q_text if q_text else q['statement']
			q['choices'] = choices
		else:
			# try a different pattern: lines starting with A) B) etc
			line_pattern = re.compile(
				r'^\s*([A-E])[)]\s*(.+?)$', re.MULTILINE
			)
			line_matches = list(line_pattern.finditer(text))
			if len(line_matches) >= 2:
				# everything before first match is question continuation
				pre_text = text[:line_matches[0].start()].strip()
				if pre_text:
					q['statement'] = pre_text
				choices = []
				for m in line_matches:
					choice_text = m.group(2).strip()
					choices.append(choice_text)
				q['choices'] = choices

	return questions


#============================================
def build_yaml_structure(doc: docx.Document, questions: list) -> dict:
	"""Build the full YAML exam structure.

	Args:
		doc: The loaded DOCX document.
		questions: List of parsed question dicts.

	Returns:
		Dict matching the odt_exam_builder YAML format.
		Uses 'chapter' key for section headings instead of 'heading'.
	"""
	# extract header info from first few paragraphs
	paras = doc.paragraphs
	student_line = paras[0].text.strip() if len(paras) > 0 else ''
	title = paras[1].text.strip() if len(paras) > 1 else 'Exam'
	# build sections - group all questions into one section for now
	section = {
		'chapter': 'Multiple Choice',
		'questions': questions,
	}
	exam_data = {
		'title': title,
		'student_line': student_line,
		'sections': [section],
	}
	return exam_data


#============================================
def main():
	"""Main entry point for DOCX to YAML converter."""
	args = parse_args()
	# load the DOCX
	doc = docx.Document(args.input_file)
	# extract images
	para_images = extract_images(doc, args.image_dir)
	print(f"Extracted {len(para_images)} images to {args.image_dir}/")
	# parse questions
	questions = parse_docx_questions(doc, para_images)
	print(f"Parsed {len(questions)} questions")
	# build YAML structure
	exam_data = build_yaml_structure(doc, questions)
	# write YAML
	with open(args.output_file, 'w') as f:
		yaml.dump(exam_data, f, default_flow_style=False, allow_unicode=True, width=120)
	print(f"YAML written to {args.output_file}")


#============================================
if __name__ == '__main__':
	main()
