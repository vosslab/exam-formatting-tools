#!/usr/bin/env python3
"""Generate a styled DOCX exam document from structured YAML input.

Uses python-docx to create a Word document with exam styles, auto-numbering,
bold choice prefixes, images, tables, and page headers. Reads the YAML exam
format (see docs/YAML_EXAM_FORMAT.md) and styles from styles/exam_styles.yaml.
"""

# Standard Library
import os
import sys
import argparse

# Pip Modules
import yaml
import docx
import docx.shared
import docx.enum.text

# Local Repo Modules
import ef_tools.layout
import ef_tools.zip_grade
import ef_tools.text_utils
import ef_tools.cli_checks
import ef_tools.style_loader
import ef_tools.exam_defaults
import ef_tools.question_utils
import ef_tools.docx_builder


#============================================
def parse_args() -> argparse.Namespace:
	"""Parse command-line arguments.

	Returns:
		Parsed argument namespace.
	"""
	parser = argparse.ArgumentParser(
		description="Generate a styled DOCX exam document from YAML input"
	)
	parser.add_argument(
		'-i', '--input', dest='input_file', required=True,
		help="Input YAML file with exam data"
	)
	parser.add_argument(
		'-o', '--output', dest='output_file', default=None,
		help="Output DOCX file path. Defaults to <input-stem>.docx in CWD."
	)
	parser.add_argument(
		'-z', '--zip-grade', dest='zip_grade', action='store_true',
		help="ZipGrade mode: exclude nonconforming questions from DOCX output."
	)
	args = parser.parse_args()
	return args


#============================================
def apply_zip_grade_filter(exam_data: dict, source_path: str) -> dict:
	"""Filter exam_data for ZipGrade compatibility before DOCX build.

	Prints the mode banner, the per-question removal report (line-anchored),
	and a summary line showing post-filter row count. Emits a stderr WARNING
	when the post-filter row total still exceeds the ZipGrade form capacity
	(100 rows), but still returns the filtered exam so the build proceeds.

	Drops both ERROR and FIXABLE questions per the contract: a 6-choice
	question is not safe to truncate automatically. Use the
	validate_zipgrade_yaml.py linter to surface FIXABLE candidates and
	hand-edit them before re-running without the flag.

	Args:
		exam_data: Parsed exam YAML.
		source_path: Path to the input YAML file (used as report prefix).

	Returns:
		Filtered exam dict with non-OK questions removed.
	"""
	# mode banner so authors know -z is mutating output
	print("ZipGrade mode enabled: excluding nonconforming questions from DOCX output.")
	filtered, issues = ef_tools.zip_grade.filter_exam(exam_data)
	# per-question removal report (line-anchored)
	if issues:
		report_text = ef_tools.zip_grade.format_report(issues, source_path=source_path)
		print(report_text)
	# count what was dropped (per-question issues only; whole-exam issue
	# does not represent a removed question)
	dropped = sum(1 for issue in issues if issue.question_number is not None)
	kept_rows = ef_tools.zip_grade.filtered_total_rows(filtered)
	print(f"ZipGrade filter: dropped {dropped} questions; {kept_rows} rows -> DOCX")
	# warn but build when post-filter overflow remains
	if kept_rows > ef_tools.zip_grade.MAX_TOTAL_ROWS:
		warning = (f"WARNING: {kept_rows} rows after filtering; "
			f"rows {ef_tools.zip_grade.MAX_TOTAL_ROWS + 1}-{kept_rows} "
			f"won't fit on a {ef_tools.zip_grade.MAX_TOTAL_ROWS}-row ZipGrade form.")
		print(warning, file=sys.stderr)
	return filtered


#============================================
def resolve_choice_layout(question: dict, items: list, layout_limits) -> tuple:
	"""Pick (tab_style, items_per_row) for a list of choice items.

	An integer `layout` on the question bypasses auto-sizing; otherwise
	the items are sized via the shared MC auto-layout. Used by both MC
	`choices` and matching `choices_list` so the two paths stay in sync.
	"""
	layout_override = question.get('layout', None)
	if layout_override is not None:
		return layout_override, layout_override
	tab_style, items_per_row = ef_tools.layout.auto_layout_for_choices(
		items, layout_limits=layout_limits)
	return tab_style, items_per_row


#============================================
def build_document(exam_data: dict, output_path: str) -> int:
	"""Build a complete DOCX exam document from YAML data.

	Creates all styles, page layout, headers, and body content.
	Refuses to overwrite an existing file.

	Args:
		exam_data: Dict from YAML input with exam structure.
		output_path: Path for the output DOCX file.
	"""
	# enforce .docx extension
	if not output_path.endswith('.docx'):
		raise ValueError(f"Output file must have .docx extension, got: {output_path}")
	# safety check: never overwrite existing files
	if os.path.exists(output_path):
		raise FileExistsError(f"Output file already exists: {output_path}")

	doc = docx.Document()

	# load style definitions from YAML
	styles = ef_tools.style_loader.load_styles()

	# setup styles from YAML definitions
	ef_tools.docx_builder.setup_styles(doc, styles)

	# page layout: margins from YAML
	page_margin = styles['page']['margin']
	section = doc.sections[0]
	section.top_margin = docx.shared.Inches(page_margin)
	section.bottom_margin = docx.shared.Inches(page_margin)
	section.left_margin = docx.shared.Inches(page_margin)
	section.right_margin = docx.shared.Inches(page_margin)

	# setup page header (empty on first page, name/date on body pages)
	date_str = exam_data.get('date', '')
	ef_tools.docx_builder.setup_header(doc, section, date_str, styles)

	# --- first page boilerplate ---
	boilerplate_tab = styles['page']['boilerplate_tab']

	# exam title (use add_paragraph with Heading 1 style, not add_heading,
	# so our custom font overrides Word's built-in theme font)
	title = exam_data.get('title', 'Exam')
	title_para = doc.add_paragraph()
	title_para.style = doc.styles['Heading 1']
	ef_tools.docx_builder.add_rich_text_runs(title_para, title)

	# name line shifted right with tab
	name_line = exam_data.get('student_line', ef_tools.exam_defaults.DEFAULT_NAME_LINE)
	name_para = doc.add_paragraph()
	name_para.add_run("\t" + name_line)
	name_para.paragraph_format.tab_stops.add_tab_stop(
		docx.shared.Inches(boilerplate_tab), docx.enum.text.WD_TAB_ALIGNMENT.LEFT
	)

	# score line shifted right with tab
	sections = exam_data.get('sections', [])
	total_points = exam_data.get('total_points', None)
	if total_points is None:
		total_points = ef_tools.question_utils.count_total_questions(sections)
	num_sections = exam_data.get('scoring_sections', ef_tools.exam_defaults.DEFAULT_SCORING_SECTIONS)
	score_line = ef_tools.exam_defaults.format_score_line(total_points, num_sections)
	score_para = doc.add_paragraph()
	score_para.add_run("\t" + score_line)
	score_para.paragraph_format.tab_stops.add_tab_stop(
		docx.shared.Inches(boilerplate_tab), docx.enum.text.WD_TAB_ALIGNMENT.LEFT
	)

	# load image max width from styles
	image_max_width = styles['page']['image_max_width']
	# load layout limits for choice auto-sizing
	layout_limits = styles.get('layout_limits', None)
	# load style flags for table header alignment
	flags = styles['style_flags']

	# --- sections and questions ---
	question_counter = 1
	question_total = 0
	prev_element = 'chapter'
	for section_data in sections:
		# major section heading (Heading 2 level)
		heading = section_data.get('heading', '')
		if heading:
			para = doc.add_paragraph()
			para.style = doc.styles['Exam Heading 2']
			ef_tools.docx_builder.add_rich_text_runs(para, heading)
			prev_element = 'chapter'
		# chapter heading (purple, level 3)
		chapter = section_data.get('chapter', '')
		if chapter:
			para = doc.add_paragraph()
			para.style = doc.styles['Chapter Heading']
			ef_tools.docx_builder.add_rich_text_runs(para, chapter)
			prev_element = 'chapter'
		# questions
		questions = section_data.get('questions', [])
		for question in questions:
			# inline chapter heading: a YAML question-list entry of the form
			# `- chapter: "..."` is not a real question; render it as a
			# Chapter Heading paragraph and skip the question pipeline so
			# the question counter does not advance.
			if 'chapter' in question and 'statement' not in question:
				chapter_text = question['chapter']
				para = doc.add_paragraph()
				para.style = doc.styles['Chapter Heading']
				ef_tools.docx_builder.add_rich_text_runs(para, chapter_text)
				prev_element = 'chapter'
				continue
			# Hard-switch guard: matching schema migrated from `matching_terms`
			# to `prompts_list`/`choices_list`. Refuse legacy YAML loudly so
			# stale inputs don't render as malformed exams with empty prompts.
			if 'matching_terms' in question:
				raise ValueError(
					"Question contains legacy 'matching_terms' key; "
					"replace with 'prompts_list' and 'choices_list' "
					"(see docs/YAML_EXAM_FORMAT.md)."
				)
			# get explicit number if provided, otherwise use counter
			if 'number' in question:
				question_counter = question['number']
			question_number = question_counter
			# matching items use bptools-style prompts_list/choices_list:
			# prompts_list drives the question-number span and the numbered
			# blanks; choices_list is the lettered (A)/(B)/... options.
			prompts_list = question.get('prompts_list', [])
			question_span = max(1, len(prompts_list))
			question_total += question_span
			if question_span > 1:
				question_prefix = f"Q{question_number}-{question_number + question_span - 1}. "
			else:
				question_prefix = f"{question_number}. "
			# build question statement with number prefix
			statement = question.get('statement', '')
			if statement:
				# strip existing number prefix and add our own
				stripped = ef_tools.text_utils.strip_number_prefix(statement)
				# select style based on previous element
				style_name = ef_tools.question_utils.select_question_style(prev_element)
				para = doc.add_paragraph()
				para.style = doc.styles[style_name]
				# add number prefix (bold+italic inherited from style)
				para.add_run(question_prefix)
				# add statement text with rich text support
				ef_tools.docx_builder.add_rich_text_runs(para, stripped)
				prev_element = 'question'
			# matching layout matches reference artifacts in ARTIFACTS/:
			# lettered (A)/(B)/... choices come FIRST as the answer key,
			# then numbered blanks the student fills in.
			choices_list = question.get('choices_list', [])
			if choices_list:
				tab_style, items_per_row = resolve_choice_layout(
					question, choices_list, layout_limits)
				ef_tools.docx_builder.add_choices_paragraph(
					doc, choices_list, tab_style, items_per_row,
					image_width=styles['page']['choice_image_max_width'],
					image_height=styles['page']['choice_image_max_height'])
				prev_element = 'choices'
			if prompts_list:
				for index, prompt in enumerate(prompts_list):
					prompt_para = doc.add_paragraph()
					prompt_para.style = doc.styles['Matching Prompt']
					prompt_para.add_run(f"___ {question_number + index}. ")
					ef_tools.docx_builder.add_rich_text_runs(prompt_para, prompt)
				prev_element = 'choices'
			# images (before choices, after question text)
			image_paths = []
			image_path = question.get('image', None)
			if image_path is not None:
				image_paths.append(image_path)
			image_paths.extend(question.get('images', []))
			for current_image_path in image_paths:
				if os.path.isfile(current_image_path):
					# add image with auto aspect ratio, max width from styles
					doc.add_picture(current_image_path, width=docx.shared.Inches(image_max_width))
					prev_element = 'image'
			# table
			table_data = question.get('table', None)
			if table_data is not None:
				columns = table_data['columns']
				rows = table_data['rows']
				# pass table header background color and alignment from styles
				table_bg = styles['colors']['table_header_bg'].lstrip('#')
				center_header = flags['table_header_centered']
				ef_tools.docx_builder.add_table(doc, columns, rows,
					header_bg=table_bg, center_header=center_header)
				prev_element = 'table'
			# choices
			choices = question.get('choices', None)
			if choices is not None and len(choices) > 0:
				if any(isinstance(choice, dict) and choice.get('image', None) for choice in choices):
					ef_tools.docx_builder.add_image_choices_tabbed(
						doc, choices,
						image_width=styles['page']['choice_image_max_width'],
						image_height=styles['page']['choice_image_max_height'],
					)
					prev_element = 'choices'
					question_counter += question_span
					continue
				tab_style, items_per_row = resolve_choice_layout(
					question, choices, layout_limits)
				ef_tools.docx_builder.add_choices_paragraph(
					doc, choices, tab_style, items_per_row,
					image_width=styles['page']['choice_image_max_width'],
					image_height=styles['page']['choice_image_max_height'])
				prev_element = 'choices'
			# increment question counter
			question_counter += question_span

	# save document
	doc.save(output_path)
	return question_total


#============================================
def main():
	"""Main entry point for DOCX exam builder."""
	args = parse_args()
	# validate input extension and resolve default output
	ef_tools.cli_checks.require_extensions(args.input_file, ('.yml', '.yaml'), 'input')
	output_file = args.output_file
	if output_file is None:
		output_file = ef_tools.cli_checks.default_output_path(args.input_file, '.docx')
	ef_tools.cli_checks.require_extensions(output_file, ('.docx',), 'output')
	# load YAML; use the line-tracking loader when --zip-grade is set so
	# the removal report can anchor each dropped question to a source line
	if args.zip_grade:
		with open(args.input_file, 'r') as f:
			# LineTrackingLoader extends yaml.SafeLoader, so this is safe
			exam_data = yaml.load(f, Loader=ef_tools.zip_grade.LineTrackingLoader)  # nosec B506
		exam_data = apply_zip_grade_filter(exam_data, args.input_file)
	else:
		with open(args.input_file, 'r') as f:
			exam_data = yaml.safe_load(f)
	# build document and report question count
	question_count = build_document(exam_data, output_file)
	print(f"Exam DOCX written to {output_file} ({question_count} questions)")


#============================================
if __name__ == '__main__':
	main()
