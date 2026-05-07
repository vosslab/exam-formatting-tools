"""DOCX rendering helpers for exam document building.

Provides style setup, rich text rendering, choice layout, table building,
and page header configuration. All formatting values come from the styles
dict (loaded from styles/exam_styles.yaml) rather than being hardcoded.
"""

# Standard Library
import re
import datetime

# Pip Modules
import docx
import docx.shared
import docx.enum.text
import docx.enum.style
import docx.oxml.ns
import docx.oxml
import PIL.Image

# Local Repo Modules
import ef_tools.layout
import ef_tools.text_utils


#============================================
def set_font_with_fallback(style, primary: str, fallback: str) -> None:
	"""Set font name on a style with a fallback font via XML.

	python-docx only supports a single font name. This sets the primary
	font and adds the fallback as hAnsi/cs font for cross-platform
	compatibility (Liberation Sans on Linux, Arial on Windows/Mac).

	Args:
		style: A python-docx paragraph or character style.
		primary: Primary font name (e.g., 'Liberation Sans').
		fallback: Fallback font name (e.g., 'Arial').
	"""
	style.font.name = primary
	# set the fallback font on the underlying XML for non-ascii/hAnsi
	rpr = style.element.get_or_add_rPr()
	rfonts_tag = docx.oxml.ns.qn('w:rFonts')
	rfonts = rpr.find(rfonts_tag)
	if rfonts is None:
		rfonts = docx.oxml.OxmlElement('w:rFonts')
		rpr.insert(0, rfonts)
	# remove theme font attributes that override explicit font names
	# (built-in styles like Heading 1 use asciiTheme/hAnsiTheme for Calibri)
	for theme_attr in ('w:asciiTheme', 'w:hAnsiTheme', 'w:csTheme', 'w:eastAsiaTheme'):
		qname = docx.oxml.ns.qn(theme_attr)
		if rfonts.get(qname) is not None:
			del rfonts.attrib[qname]
	rfonts.set(docx.oxml.ns.qn('w:ascii'), primary)
	rfonts.set(docx.oxml.ns.qn('w:hAnsi'), fallback)
	rfonts.set(docx.oxml.ns.qn('w:cs'), fallback)


#============================================
def parse_hex_color(hex_str: str) -> docx.shared.RGBColor:
	"""Parse a hex color string like '#6600CC' into an RGBColor.

	Args:
		hex_str: Color string with leading '#'.

	Returns:
		RGBColor object.
	"""
	hex_str = hex_str.lstrip('#')
	r = int(hex_str[0:2], 16)
	g = int(hex_str[2:4], 16)
	b = int(hex_str[4:6], 16)
	color = docx.shared.RGBColor(r, g, b)
	return color


#============================================
def setup_styles(doc: docx.Document, styles: dict) -> None:
	"""Create all named exam styles in the document.

	Reads font, size, color, spacing, and style_flags values from
	the styles dict (loaded from styles/exam_styles.yaml).

	Args:
		doc: The Document to add styles to.
		styles: Style definitions dict from YAML.
	"""
	fonts = styles['fonts']
	sizes = styles['sizes']
	colors = styles['colors']
	spacing = styles['spacing']
	flags = styles['style_flags']

	# modify the built-in Normal style as our base
	normal = doc.styles['Normal']
	normal.font.size = docx.shared.Pt(sizes['normal'])
	normal.paragraph_format.space_after = docx.shared.Pt(spacing['normal_space_after_pt'])
	normal.paragraph_format.space_before = docx.shared.Pt(0)
	normal.paragraph_format.line_spacing = spacing['normal_line_spacing']
	set_font_with_fallback(normal, fonts['primary'], fonts['fallback'])

	# customize built-in Heading 1
	h1 = doc.styles['Heading 1']
	h1.font.size = docx.shared.Pt(sizes['heading_1'])
	h1.font.bold = flags['heading_1_bold']
	h1.font.italic = flags['heading_1_italic']
	h1.font.color.rgb = parse_hex_color(colors['heading_1'])
	h1.paragraph_format.space_before = docx.shared.Inches(spacing['heading_space_before'])
	h1.paragraph_format.space_after = docx.shared.Inches(spacing['heading_space_after'])
	h1.paragraph_format.keep_with_next = True
	if flags['heading_1_centered']:
		h1.paragraph_format.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
	set_font_with_fallback(h1, fonts['primary'], fonts['fallback'])

	# Question Heading: hanging indent, keep-with-next
	qh = doc.styles.add_style('Question Heading', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
	qh.base_style = normal
	qh.font.bold = flags['question_bold']
	qh.font.italic = flags['question_italic']
	qh.font.size = docx.shared.Pt(sizes['question'])
	qh.paragraph_format.left_indent = docx.shared.Inches(spacing['question_indent'])
	qh.paragraph_format.first_line_indent = docx.shared.Inches(spacing['question_hanging'])
	qh.paragraph_format.space_before = docx.shared.Inches(spacing['question_heading_space_before'])
	qh.paragraph_format.space_after = docx.shared.Inches(spacing['question_space_after'])
	qh.paragraph_format.keep_with_next = True

	# Question Follow: inherits Question Heading. The artifact (e.g.
	# ARTIFACTS/exam1.docx) only overrides space_before=0 on top of
	# Question Heading; everything else (bold, italic, 11pt, indent,
	# hanging, keep_with_next) is inherited. Mirror that here so changes
	# to Question Heading propagate correctly.
	qf = doc.styles.add_style('Question Follow', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
	qf.base_style = qh
	qf.paragraph_format.space_before = docx.shared.Pt(0)
	qf.paragraph_format.space_after = docx.shared.Inches(spacing['question_space_after'])

	# Chapter Heading: colored, keep-with-next
	ch = doc.styles.add_style('Chapter Heading', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
	ch.base_style = normal
	ch.font.size = docx.shared.Pt(sizes['chapter_heading'])
	set_font_with_fallback(ch, fonts['primary'], fonts['fallback'])
	ch.font.bold = flags['chapter_heading_bold']
	ch.font.color.rgb = parse_hex_color(colors['chapter_heading'])
	ch.paragraph_format.space_before = docx.shared.Pt(spacing['chapter_space_before_pt'])
	ch.paragraph_format.space_after = docx.shared.Pt(spacing['chapter_space_after_pt'])
	ch.paragraph_format.keep_with_next = True

	# Heading 2: for major section labels
	h2 = doc.styles.add_style('Exam Heading 2', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
	h2.base_style = normal
	h2.font.size = docx.shared.Pt(sizes['heading_2'])
	set_font_with_fallback(h2, fonts['primary'], fonts['fallback'])
	h2.font.bold = flags['heading_2_bold']
	h2.font.italic = flags['heading_2_italic']
	h2.paragraph_format.space_before = docx.shared.Pt(spacing['chapter_space_before_pt'])
	h2.paragraph_format.space_after = docx.shared.Pt(spacing['chapter_space_after_pt'])
	h2.paragraph_format.keep_with_next = True

	# Choice: base style for all choice layouts
	choice_base = doc.styles.add_style('Choice', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
	choice_base.base_style = normal
	choice_base.font.size = docx.shared.Pt(sizes['choice'])
	choice_base.paragraph_format.left_indent = docx.shared.Inches(spacing['choice_indent'])
	choice_first_line = spacing['choice_first_line_indent']
	if choice_first_line != 0:
		choice_base.paragraph_format.first_line_indent = docx.shared.Inches(choice_first_line)
	choice_base.paragraph_format.space_before = docx.shared.Pt(0)
	choice_base.paragraph_format.space_after = docx.shared.Pt(spacing['normal_space_after_pt'])

	# Header: small font for page headers with center and right tab stops
	# modify the built-in Header style (not add_style)
	hdr = doc.styles['Header']
	hdr.font.size = docx.shared.Pt(sizes['header'])
	# use header-specific fonts if defined, otherwise fall back to primary
	header_primary = fonts.get('header_primary', fonts['primary'])
	header_fallback = fonts.get('header_fallback', fonts['fallback'])
	set_font_with_fallback(hdr, header_primary, header_fallback)
	# small gap below header before content
	hdr.paragraph_format.space_after = docx.shared.Pt(spacing['header_space_after_pt'])
	# clear any pre-existing tab stops from the built-in Header style
	hdr.paragraph_format.tab_stops.clear_all()
	# left-aligned page number, center date, right-aligned name
	hdr.paragraph_format.tab_stops.add_tab_stop(
		docx.shared.Inches(styles['page']['header_center_tab']),
		docx.enum.text.WD_TAB_ALIGNMENT.CENTER
	)
	hdr.paragraph_format.tab_stops.add_tab_stop(
		docx.shared.Inches(styles['page']['header_right_tab']),
		docx.enum.text.WD_TAB_ALIGNMENT.RIGHT
	)

	# Matching Prompt: numbered fill-in lines for matching questions.
	# Mirrors the legacy ARTIFACTS/2019_exam2-final.docx "Question" style
	# (verified via LibreOffice style dialog). Two tab stops let prompts
	# render in a two-column layout: "___ 1. text<tab>___ 2. text".
	matching_prompt = doc.styles.add_style(
		'Matching Prompt', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
	matching_prompt.base_style = qh
	matching_prompt.font.bold = False
	matching_prompt.font.italic = False
	matching_prompt.paragraph_format.left_indent = docx.shared.Inches(spacing['matching_prompt_indent'])
	matching_prompt.paragraph_format.first_line_indent = docx.shared.Inches(spacing['matching_prompt_hanging'])
	matching_prompt.paragraph_format.space_before = docx.shared.Inches(spacing['matching_prompt_space_before'])
	matching_prompt.paragraph_format.space_after = docx.shared.Pt(0)
	matching_prompt.paragraph_format.line_spacing = spacing['matching_prompt_line_spacing']
	matching_prompt.paragraph_format.keep_with_next = False
	for pos in styles['matching_prompt_tab_stops']:
		matching_prompt.paragraph_format.tab_stops.add_tab_stop(docx.shared.Inches(pos))

	# Choices 2 through 5: inherit from Choice, with tab stops on the style
	# load tab stops from YAML if available, else use layout module defaults
	tab_stops_config = styles.get('layout_tab_stops', ef_tools.layout.CHOICES_TAB_STOPS)
	for n in range(2, 6):
		style = doc.styles.add_style(f'Choices {n}', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
		style.base_style = choice_base
		# add tab stops so switching styles in Word preserves column alignment
		stops = tab_stops_config.get(n, tab_stops_config.get(str(n), []))
		for pos in stops:
			style.paragraph_format.tab_stops.add_tab_stop(docx.shared.Inches(pos))


#============================================
def add_page_number_field(paragraph) -> None:
	"""Add 'Page X of Y' field codes to a paragraph using raw XML.

	python-docx does not have native page number field support,
	so we insert the Word XML field elements directly.

	Args:
		paragraph: The paragraph to add page numbers to.
	"""
	# "Page " text (font size inherited from paragraph style)
	paragraph.add_run("Page ")
	# PAGE field
	run1 = paragraph.add_run()
	fld_simple_page = docx.oxml.OxmlElement('w:fldSimple')
	fld_simple_page.set(docx.oxml.ns.qn('w:instr'), ' PAGE ')
	run1._element.addnext(fld_simple_page)
	# " of " text
	paragraph.add_run(" of ")
	# NUMPAGES field
	run2 = paragraph.add_run()
	fld_simple_total = docx.oxml.OxmlElement('w:fldSimple')
	fld_simple_total.set(docx.oxml.ns.qn('w:instr'), ' NUMPAGES ')
	run2._element.addnext(fld_simple_total)


#============================================
def format_date_human(date_str: str) -> str:
	"""Convert an ISO date string to human-readable format.

	Converts '2026-04-02' to 'April 2, 2026'. If parsing fails,
	returns the original string unchanged.

	Args:
		date_str: Date string in ISO format (YYYY-MM-DD).

	Returns:
		Human-readable date string like 'April 2, 2026'.
	"""
	# parse ISO date and format with full month name, no zero-padded day
	dt = datetime.datetime.strptime(date_str, '%Y-%m-%d')
	# %-d gives non-padded day on Unix/macOS
	formatted = dt.strftime('%B %-d, %Y')
	return formatted

assert format_date_human('2026-04-02') == 'April 2, 2026'
assert format_date_human('2025-12-25') == 'December 25, 2025'
assert format_date_human('2026-01-01') == 'January 1, 2026'


#============================================
def setup_header(doc: docx.Document, section, date_str: str, styles: dict) -> None:
	"""Configure page header on body pages with page number, date, and name.

	Sets up different first page header (empty) and body page header with
	'Page X of Y | date | First Name:___' layout. Font size and tab stops
	come from the Header style.

	Args:
		doc: The Document (needed to access Header style).
		section: The document section to configure.
		date_str: Date string to display in header center.
		styles: Style definitions dict from YAML.
	"""
	# enable different first page header (empty on page 1)
	section.different_first_page_header_footer = True
	# fix header distance so top margin is consistent on all pages
	header_distance = styles['page']['header_distance']
	section.header_distance = docx.shared.Inches(header_distance)

	# body page header: "Page X of Y [tab] date [tab] First Name:___"
	header = section.header
	header.is_linked_to_previous = False
	header_para = header.paragraphs[0]
	# apply Header style (font size and tab stops defined on style)
	header_para.style = doc.styles['Header']
	# add page number fields (left-aligned by default)
	add_page_number_field(header_para)
	# tab to center, add human-readable date
	header_para.add_run("\t")
	if date_str:
		# convert ISO date (2026-04-02) to human-readable (April 2, 2026)
		formatted_date = format_date_human(date_str)
		header_para.add_run(formatted_date)
	# tab to right, add name line
	header_para.add_run("\t")
	header_para.add_run("First Name:_____________________________")

	# first page header is empty (just clear it)
	first_header = section.first_page_header
	# clear any default content
	for p in first_header.paragraphs:
		p.text = ""


#============================================
def add_rich_text_runs(para, text: str) -> None:
	"""Add styled runs to a paragraph, parsing inline HTML tags.

	Handles <sub>, <sup>, <b>, <strong>, <i>, <em> tags by creating
	separate runs with appropriate font properties. HTML entities are
	decoded first, then rich text tags are parsed.

	Font size, base bold, and base italic come from the paragraph style.
	Only rich text tags add run-level overrides.

	Args:
		para: The paragraph to add runs to.
		text: Text that may contain HTML entities and inline tags.
	"""
	# decode HTML entities first, then parse rich text tags
	decoded = ef_tools.text_utils.decode_html_entities(text)
	segments = ef_tools.text_utils.parse_rich_text(decoded)
	for segment_text, tags in segments:
		# drop hard-break segments and collapse stray newlines to spaces;
		# paragraph-level splitting belongs to add_rich_text_paragraphs
		if segment_text == "\n":
			continue
		flat_text = segment_text.replace("\n", " ")
		if not flat_text:
			continue
		run = para.add_run(flat_text)
		# only set run-level overrides for rich text tags
		if 'b' in tags:
			run.font.bold = True
		if 'i' in tags:
			run.font.italic = True
		if 'sub' in tags:
			run.font.subscript = True
		if 'sup' in tags:
			run.font.superscript = True


# regex to split text on hard paragraph breaks: literal \n or <br> tags
_PARAGRAPH_BREAK_RE = re.compile(r'(?:\r?\n|<br\s*/?>)+', re.IGNORECASE)


#============================================
def add_rich_text_paragraphs(doc, style_name: str, text: str,
		prefix: str = '') -> object:
	"""Split text on hard breaks (\\n and <br>) into separate paragraphs.

	Each piece becomes its own paragraph styled with style_name. The
	optional prefix is added as a plain run on the first paragraph only.
	An empty input still emits one paragraph (carrying just the prefix)
	so callers can rely on at least one paragraph being created.

	Returns the last paragraph created.
	"""
	pieces = _PARAGRAPH_BREAK_RE.split(text)
	# keep only pieces with visible content; preserve at least one slot
	pieces = [piece for piece in pieces if piece.strip()]
	if not pieces:
		pieces = ['']
	last_para = None
	for index, piece in enumerate(pieces):
		para = doc.add_paragraph()
		para.style = doc.styles[style_name]
		if index == 0 and prefix:
			para.add_run(prefix)
		add_rich_text_runs(para, piece)
		last_para = para
	return last_para


#============================================
def fit_picture_kwargs(image_path: str, max_width: float,
	max_height: float = None) -> dict:
	"""Pick add_picture(width=...) or (height=...) so the image fits a box.

	python-docx preserves aspect ratio when only one of width/height is
	set. To fit a bounding box (max_width, max_height), pick the axis
	whose scale factor is smaller -- that axis becomes the binding
	constraint. Returns a kwargs dict for add_picture.
	"""
	# no height bound: fall back to width-only behavior
	if max_height is None:
		return {'width': docx.shared.Inches(max_width)}
	with PIL.Image.open(image_path) as img:
		src_w, src_h = img.size
	# scale factors needed to hit each bound; the smaller one wins
	width_scale = max_width / src_w
	height_scale = max_height / src_h
	if width_scale <= height_scale:
		return {'width': docx.shared.Inches(max_width)}
	return {'height': docx.shared.Inches(max_height)}


#============================================
def _zero_inline_image_margins(run) -> None:
	"""Set distT/distB/distL/distR=0 on every <wp:inline> in this run.

	Inline images in OOXML carry top/bottom/left/right distance attributes
	that act as outer padding around the picture. python-docx's add_picture
	leaves them at default values (often 114300 EMU = 0.125"), which chews
	visible space around each picture in a tight tab-separated row. Zeroing
	them lets the image hug its run's edges so the per-column budget can
	hold a wider visible image at the same docx column width.
	"""
	WP = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
	for inline in run.element.iter('{%s}inline' % WP):
		for attr in ('distT', 'distB', 'distL', 'distR'):
			inline.set(attr, '0')


#============================================
def add_choice_content(para, choice, image_width: float = None,
	image_height: float = None) -> None:
	"""Add text and optional image content for one choice."""
	choice_text = ef_tools.layout.choice_text(choice)
	if choice_text:
		add_rich_text_runs(para, choice_text)
	image_path = ef_tools.layout.choice_image(choice)
	if image_path:
		if choice_text:
			para.add_run().add_break()
		run = para.add_run()
		if image_width is None:
			run.add_picture(image_path)
		else:
			kwargs = fit_picture_kwargs(image_path, image_width, image_height)
			run.add_picture(image_path, **kwargs)


#============================================
# Maximum image width per column count, in inches. Empirical caps tuned via
# _render_loop.sh / _measure.py with the Choices N tab stops in
# styles/exam_styles.yaml. Going past these widths pushes a trailing image's
# cursor past its target tab stop, which then advances to the NEXT stop and
# creates a visible gap. Word's inline-image layout reserves more space per
# image than a naive cursor calculation predicts; values are set by viewing
# the rendered DOCX, not computed from page width.
IMAGE_CHOICE_MAX_WIDTH_BY_COLS = {
	2: 3.30,
	3: 2.07,
	4: 1.49,
	5: 1.13,
}


#============================================
def _is_meaningful_alt(choice) -> bool:
	"""True when a choice's text adds information beyond a placeholder.

	Skips empty strings and the literal "image" so a row of placeholder
	alt text does not render as a redundant caption line. Used to decide
	whether an alt-text paragraph is emitted under an image-choice row.
	"""
	if not ef_tools.layout.choice_image(choice):
		return False
	text = (ef_tools.layout.choice_text(choice) or '').strip().lower()
	if not text or text == 'image':
		return False
	return True


#============================================
def add_image_choices_tabbed(doc: docx.Document, choices: list,
	image_width: float, image_height: float = None) -> None:
	"""Add image-based choices in a horizontal tab-stop layout.

	Letter prefix and image render on the same line for each column.
	When at least one choice carries meaningful alt text (per
	`_is_meaningful_alt`), a second paragraph below holds the alt-text
	row. Tab stops are inherited from the Choices N paragraph style; no
	paragraph-level tab stops are added (style stops at the same column
	positions cause Word to see two close-but-not-equal stops per column
	and images drift off-grid). No docx tables are created.

	Args:
		doc: The Document to add the paragraph to.
		choices: List of structured choice dicts with text and/or image keys.
		image_width: Maximum image width in inches from the YAML style
			cap; further clamped by IMAGE_CHOICE_MAX_WIDTH_BY_COLS so 5-col
			rows stay narrow enough not to wrap.
		image_height: Maximum image height in inches; aspect ratio preserved.
	"""
	# clamp column count to the legal Choices 2..5 range so the
	# resolved style is always concrete (never the abstract Choice base)
	num_cols = len(choices)
	style_columns = max(2, min(num_cols, 5))
	style_name = ef_tools.layout.choices_style_name(style_columns)
	# image width: the YAML-supplied cap (image_width) plus a per-num_cols
	# empirical cap (IMAGE_CHOICE_MAX_WIDTH_BY_COLS); the smaller wins.
	# Column counts beyond 5 (matching questions with many panels) fall
	# back to the 5-col cap (the most conservative empirical value).
	per_cols_cap = IMAGE_CHOICE_MAX_WIDTH_BY_COLS.get(num_cols,
		IMAGE_CHOICE_MAX_WIDTH_BY_COLS[5])
	effective_image_width = min(image_width, per_cols_cap)

	# Single combined paragraph: (A)<img> \t (B)<img> \t (C)<img> \t ...
	# Letter and image render side-by-side at each tab stop. Inherits
	# left_indent from the Choices N style so image-choice rows align
	# vertically with text-only Choices rows above/below them. Tab stops
	# in styles/exam_styles.yaml -> layout_tab_stops are pre-offset by
	# choice_indent so all inter-column gaps render evenly.
	combined_para = doc.add_paragraph()
	combined_para.style = doc.styles[style_name]
	combined_para.paragraph_format.keep_with_next = True
	combined_para.paragraph_format.space_after = docx.shared.Pt(0)
	# unify height across the row: pick the smallest height each image would
	# have if scaled to fit the (effective_image_width, image_height) box,
	# then force every image in the row to that exact height. Keeps the row
	# visually level even when source PNGs have different aspect ratios.
	row_height = _row_image_height(choices, effective_image_width, image_height)
	for index, choice in enumerate(choices):
		if index > 0:
			combined_para.add_run("\t")
		letter = chr(ord('A') + index)
		image_path = ef_tools.layout.choice_image(choice)
		# Image-bearing choices: bold "(A)" sits on the same line as the
		# image, no trailing space. Text-only choices: bold "(A) " plus
		# caption text on the same line, separator space.
		prefix_text = f"({letter})" if image_path else f"({letter}) "
		prefix = combined_para.add_run(prefix_text)
		prefix.bold = True
		if image_path:
			image_run = combined_para.add_run()
			if row_height is not None:
				image_run.add_picture(image_path,
					height=docx.shared.Inches(row_height))
			else:
				kwargs = fit_picture_kwargs(image_path,
					effective_image_width, image_height)
				image_run.add_picture(image_path, **kwargs)
			# zero the inline image's edge-distance margins so the image
			# hugs the run; prevents the default ~0.125" padding that
			# python-docx injects around inline pictures.
			_zero_inline_image_margins(image_run)
		else:
			choice_text = ef_tools.layout.choice_text(choice)
			if choice_text:
				add_rich_text_runs(combined_para, choice_text)

	# alt-text paragraph: rendered only when at least one image-bearing
	# choice carries non-trivial text. Skipping the row when every alt is
	# just "image" or empty avoids a redundant line of placeholder text.
	if any(_is_meaningful_alt(choice) for choice in choices):
		alt_para = doc.add_paragraph()
		alt_para.style = doc.styles[style_name]
		alt_para.paragraph_format.space_before = docx.shared.Pt(0)
		for index, choice in enumerate(choices):
			if index > 0:
				alt_para.add_run("\t")
			if ef_tools.layout.choice_image(choice):
				choice_text = ef_tools.layout.choice_text(choice)
				if choice_text:
					add_rich_text_runs(alt_para, choice_text)


#============================================
def _row_image_height(choices: list, max_width: float,
	max_height: float = None) -> float:
	"""Single row-level height that lets every image fit its column box.

	For each image, compute the height it would be when scaled to fit
	inside (max_width, max_height) preserving aspect. Return the minimum,
	so the binding constraint of the row sets a uniform height. Returns
	None if no image paths are present or no height bound is given (then
	callers fall back to per-image fit_picture_kwargs).
	"""
	if max_height is None:
		return None
	heights = []
	for choice in choices:
		image_path = ef_tools.layout.choice_image(choice)
		if not image_path:
			continue
		with PIL.Image.open(image_path) as img:
			src_w, src_h = img.size
		# height when width is the binding constraint
		h_at_max_w = max_width * src_h / src_w
		# this image's rendered height is whichever bound is tighter
		heights.append(min(h_at_max_w, max_height))
	if not heights:
		return None
	return min(heights)


#============================================
def add_choices_paragraph(doc: docx.Document, choices: list,
	tab_style: int, items_per_row: int, image_width: float = None,
	image_height: float = None) -> None:
	"""Add a tab-separated choices paragraph with bold letter prefixes.

	Creates a paragraph with (A) (B) (C) format, using tab characters
	to align columns. The tab_style selects tab stop positions (3, 4, 5)
	and items_per_row controls how many choices per line.

	Args:
		doc: The Document to add the paragraph to.
		choices: List of choice text strings or dicts with text/image keys.
		tab_style: Tab stop layout (3, 4, or 5) from CHOICES_TAB_STOPS.
		items_per_row: Number of choices per line (may differ from tab_style).
	"""
	# any image in the list forces a vertical stack so the image and
	# its caption stay on one line per choice
	has_images = any(ef_tools.layout.choice_image(choice) for choice in choices)
	if has_images:
		items_per_row = 1
		tab_style = 1
	# select the appropriate style via the central resolver, which
	# always returns a concrete Choices N (never the bare Choice base)
	style_name = ef_tools.layout.choices_style_name(tab_style)
	# group choices into rows so each row becomes its own paragraph
	# (hard breaks). items_per_row<=1 means one choice per row.
	row_size = max(1, items_per_row)
	rows = [choices[start:start + row_size]
		for start in range(0, len(choices), row_size)]
	for row_index, row_choices in enumerate(rows):
		para = doc.add_paragraph()
		para.style = doc.styles[style_name]
		# rows after the first should sit flush against the prior row;
		# all rows except the last drop trailing space so the choices
		# group reads as one unit visually
		if row_index > 0:
			para.paragraph_format.space_before = docx.shared.Pt(0)
		if row_index < len(rows) - 1:
			para.paragraph_format.space_after = docx.shared.Pt(0)
		for col_index, choice in enumerate(row_choices):
			# global letter index across all rows
			letter = chr(ord('A') + row_index * row_size + col_index)
			# tab before this choice (except first in each row)
			if col_index > 0:
				para.add_run("\t")
			# bold letter prefix (only run-level override needed)
			bold_run = para.add_run(f"({letter}) ")
			bold_run.bold = True
			# choice content inherits font size from style
			add_choice_content(para, choice,
				image_width=image_width, image_height=image_height)


#============================================
def add_table(doc: docx.Document, columns: list, rows: list,
	header_bg: str = 'F2F2F2', center_header: bool = True) -> None:
	"""Add a table with a styled header row.

	The header row uses bold text with a background color.

	Args:
		doc: The Document to add the table to.
		columns: List of column header strings.
		rows: List of row data (each row is a list of cell strings).
		header_bg: Hex color for header background (no '#' prefix).
		center_header: Whether to center-align header cells.
	"""
	num_rows = len(rows) + 1
	num_cols = len(columns)
	table = doc.add_table(rows=num_rows, cols=num_cols)
	table.style = 'Table Grid'
	# header row
	for ci, col_text in enumerate(columns):
		cell = table.rows[0].cells[ci]
		cell_para = cell.paragraphs[0]
		# add header text with rich text support (decodes HTML entities)
		add_rich_text_runs(cell_para, col_text)
		# set bold on all runs in the header cell
		for run in cell_para.runs:
			run.font.bold = True
		if center_header:
			cell_para.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
		# background color via XML shading
		shading = docx.oxml.OxmlElement('w:shd')
		shading.set(docx.oxml.ns.qn('w:fill'), header_bg)
		shading.set(docx.oxml.ns.qn('w:val'), 'clear')
		cell._element.get_or_add_tcPr().append(shading)
	# data rows
	for ri, row_data in enumerate(rows):
		for ci, cell_text in enumerate(row_data):
			cell_para = table.rows[ri + 1].cells[ci].paragraphs[0]
			add_rich_text_runs(cell_para, cell_text)
