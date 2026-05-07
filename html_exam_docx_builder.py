#!/usr/bin/env python3
"""Generate a styled DOCX exam from cleaned Blackboard HTML exports."""

# Standard Library
import argparse
import datetime
import os
import re

# Pip Modules
import docx
import docx.enum.text
import docx.shared
import lxml.html

# Local Repo Modules
import ef_tools.docx_builder
import ef_tools.exam_defaults
import ef_tools.layout
import ef_tools.question_utils
import ef_tools.style_loader


TAKE_QUESTION_XPATH = (
	"//div[contains(concat(' ', normalize-space(@class), ' '), ' takeQuestionDiv ')]"
)
IMAGE_CLASS_XPATH = (
	".//img[contains(concat(' ', normalize-space(@class), ' '), ' cleaned-statement-media ') "
	"or contains(concat(' ', normalize-space(@class), ' '), ' cleaned-choice-media ')]"
)
QUESTION_CODE_RE = re.compile(r"^[0-9a-f]{2,}(?:_[0-9a-f]{2,})*$", re.IGNORECASE)
CHOICE_PREFIX_RE = re.compile(r"^[A-Z][.)]\s+")
INLINE_TAGS = {
	"b",
	"strong",
	"i",
	"em",
	"sub",
	"sup",
}


#============================================
def parse_args() -> argparse.Namespace:
	"""Parse command-line arguments."""
	parser = argparse.ArgumentParser(
		description="Generate one styled DOCX exam from cleaned Blackboard HTML files."
	)
	parser.add_argument(
		"-i", "--input", dest="input_files", nargs="+", required=True,
		help="Input cleaned HTML file(s), in document order."
	)
	parser.add_argument(
		"-o", "--output", dest="output_file", required=True,
		help="Output DOCX file path, which must not already exist."
	)
	parser.add_argument(
		"-t", "--title", dest="title", default="Final Exam",
		help="Exam title for the first page."
	)
	parser.add_argument(
		"--date", dest="date", default=None,
		help="ISO date for page headers (defaults to today)."
	)
	args = parser.parse_args()
	if args.date is None:
		args.date = datetime.date.today().isoformat()
	return args


#============================================
def extract_document_title(path: str) -> str:
	"""Extract a readable title from an HTML document."""
	with open(path, "r", encoding="utf-8") as handle:
		text = handle.read()
	doc = lxml.html.fromstring(text)
	title = doc.xpath("string(//title)").strip()
	title = re.sub(r"\s+", " ", title)
	title = title.replace("Preview Test:", "").strip()
	if " - " in title:
		title = title.split(" - ", 1)[0].strip()
	if "-" in title:
		title = title.split("-", 1)[0].strip()
	return title


#============================================
def is_question_code(text: str) -> bool:
	"""Return whether text is a hidden Blackboard-ish question code."""
	clean_text = re.sub(r"\s+", "", text)
	result = bool(QUESTION_CODE_RE.match(clean_text))
	return result


#============================================
def strip_choice_prefix(text: str) -> str:
	"""Remove an existing A./B. choice prefix from exported choice text."""
	clean_text = re.sub(r"\s+", " ", text).strip()
	clean_text = CHOICE_PREFIX_RE.sub("", clean_text)
	return clean_text


#============================================
def resolve_image_path(html_path: str, src: str) -> str:
	"""Resolve an image src relative to its HTML file."""
	if src.startswith(("http://", "https://")):
		raise ValueError(f"Remote image sources are not supported: {src}")
	base_dir = os.path.dirname(html_path)
	image_path = os.path.normpath(os.path.join(base_dir, src))
	return image_path


#============================================
def text_from_element(element) -> str:
	"""Extract normalized text from an HTML element."""
	text = element.text_content()
	text = re.sub(r"\s+", " ", text).strip()
	return text


#============================================
def has_choice_labels(element) -> bool:
	"""Return whether an element contains exported multiple-choice labels."""
	labels = element.xpath(".//label[.//input]")
	result = len(labels) > 0
	return result


#============================================
def has_image(element) -> bool:
	"""Return whether an element contains relevant image content."""
	images = element.xpath(IMAGE_CLASS_XPATH)
	result = len(images) > 0
	return result


#============================================
def add_inline_content(paragraph, element, active_tags: set = None) -> None:
	"""Add an HTML element's inline text content to a paragraph."""
	if active_tags is None:
		active_tags = set()
	tag = element.tag.lower() if isinstance(element.tag, str) else ""
	child_tags = set(active_tags)
	if tag in INLINE_TAGS:
		if tag in ("strong", "b"):
			child_tags.add("b")
		elif tag in ("em", "i"):
			child_tags.add("i")
		else:
			child_tags.add(tag)
	if element.text:
		add_text_run(paragraph, element.text, child_tags)
	for child in element:
		child_tag = child.tag.lower() if isinstance(child.tag, str) else ""
		if child_tag == "br":
			paragraph.add_run().add_break()
		elif child_tag not in ("input", "img", "script", "style", "canvas"):
			add_inline_content(paragraph, child, child_tags)
		if child.tail:
			add_text_run(paragraph, child.tail, child_tags)


#============================================
def add_text_run(paragraph, text: str, active_tags: set) -> None:
	"""Add one styled text run to a paragraph."""
	if text == "":
		return
	run = paragraph.add_run(text)
	if "b" in active_tags:
		run.font.bold = True
	if "i" in active_tags:
		run.font.italic = True
	if "sub" in active_tags:
		run.font.subscript = True
	if "sup" in active_tags:
		run.font.superscript = True


#============================================
def add_question_paragraph(doc, question_number: int, element, style_name: str) -> None:
	"""Add the main question statement paragraph."""
	paragraph = doc.add_paragraph()
	paragraph.style = doc.styles[style_name]
	prefix = paragraph.add_run(f"{question_number}. ")
	prefix.bold = True
	add_inline_content(paragraph, element)


#============================================
def add_standard_paragraph(doc, element) -> None:
	"""Add a normal paragraph from an HTML element."""
	text = text_from_element(element)
	if not text:
		return
	if is_question_code(text):
		return
	paragraph = doc.add_paragraph()
	add_inline_content(paragraph, element)


#============================================
def add_image(doc, html_path: str, image_element, max_width: float) -> None:
	"""Add an HTML image to the DOCX document."""
	src = image_element.get("src", "")
	if not src:
		return
	image_path = resolve_image_path(html_path, src)
	if not os.path.isfile(image_path):
		raise FileNotFoundError(f"Image not found: {image_path}")
	doc.add_picture(image_path, width=docx.shared.Inches(max_width))


#============================================
def add_choice_label_paragraph(doc, html_path: str, labels: list, styles: dict) -> None:
	"""Add exported multiple-choice labels to the DOCX."""
	choice_texts = []
	all_text_only = True
	for label in labels:
		clean_text = strip_choice_prefix(text_from_element(label))
		choice_texts.append(clean_text)
		if has_image(label):
			all_text_only = False
	if all_text_only:
		layout_limits = styles.get("layout_limits", None)
		tab_style, items_per_row = ef_tools.layout.auto_layout_for_choices(
			choice_texts,
			layout_limits=layout_limits,
		)
		ef_tools.docx_builder.add_choices_paragraph(
			doc,
			choice_texts,
			tab_style,
			items_per_row,
		)
		return
	# build structured choices for tab-stop horizontal layout
	structured_choices = []
	for label in labels:
		clean_text = strip_choice_prefix(text_from_element(label))
		image_path = None
		images = label.xpath(IMAGE_CLASS_XPATH)
		if images:
			image_path = resolve_image_path(html_path, images[0].get("src", ""))
			if not os.path.isfile(image_path):
				raise FileNotFoundError(f"Image not found: {image_path}")
		structured_choices.append({"text": clean_text, "image": image_path})
	ef_tools.docx_builder.add_image_choices_tabbed(
		doc,
		structured_choices,
		image_width=styles["page"]["choice_image_max_width"],
		image_height=styles["page"]["choice_image_max_height"],
	)


#============================================
def add_matching_block(doc, element) -> None:
	"""Add a non-label div, usually matching prompts or answer blanks."""
	text = text_from_element(element)
	if not text:
		return
	paragraph = doc.add_paragraph()
	add_inline_content(paragraph, element)


#============================================
def add_question_body(doc, html_path: str, question_div, question_number: int,
	styles: dict, prev_element: str) -> str:
	"""Add one exported Blackboard question to the DOCX document."""
	li_elements = question_div.xpath("./li")
	if li_elements:
		body = li_elements[0]
	else:
		body = question_div
	style_name = ef_tools.question_utils.select_question_style(prev_element)
	first_statement = True
	image_max_width = styles["page"]["image_max_width"]
	for child in body:
		tag = child.tag.lower() if isinstance(child.tag, str) else ""
		if tag in ("script", "style", "input"):
			continue
		if tag == "h5":
			continue
		if has_choice_labels(child):
			labels = child.xpath(".//label[.//input]")
			add_choice_label_paragraph(doc, html_path, labels, styles)
			first_statement = False
			prev_element = "choices"
			continue
		if has_image(child):
			text = text_from_element(child)
			if text and not is_question_code(text):
				if first_statement:
					add_question_paragraph(doc, question_number, child, style_name)
					first_statement = False
				else:
					add_standard_paragraph(doc, child)
			for image in child.xpath(IMAGE_CLASS_XPATH):
				add_image(doc, html_path, image, image_max_width)
			prev_element = "image"
			continue
		text = text_from_element(child)
		if not text or is_question_code(text):
			continue
		if first_statement:
			add_question_paragraph(doc, question_number, child, style_name)
			first_statement = False
			prev_element = "question"
		elif tag == "div":
			add_matching_block(doc, child)
			prev_element = "table"
		else:
			add_standard_paragraph(doc, child)
			prev_element = "question"
	if first_statement:
		paragraph = doc.add_paragraph()
		paragraph.style = doc.styles[style_name]
		paragraph.add_run(f"{question_number}. ")
		prev_element = "question"
	return prev_element


#============================================
def add_front_page(doc, title: str, date: str, total_points: int, styles: dict) -> None:
	"""Add title, name line, and score line."""
	boilerplate_tab = styles["page"]["boilerplate_tab"]
	title_para = doc.add_paragraph()
	title_para.style = doc.styles["Heading 1"]
	ef_tools.docx_builder.add_rich_text_runs(title_para, title)
	name_para = doc.add_paragraph()
	name_para.add_run("\t" + ef_tools.exam_defaults.DEFAULT_NAME_LINE)
	name_para.paragraph_format.tab_stops.add_tab_stop(
		docx.shared.Inches(boilerplate_tab),
		docx.enum.text.WD_TAB_ALIGNMENT.LEFT,
	)
	score_line = ef_tools.exam_defaults.format_score_line(
		total_points,
		ef_tools.exam_defaults.DEFAULT_SCORING_SECTIONS,
	)
	score_para = doc.add_paragraph()
	score_para.add_run("\t" + score_line)
	score_para.paragraph_format.tab_stops.add_tab_stop(
		docx.shared.Inches(boilerplate_tab),
		docx.enum.text.WD_TAB_ALIGNMENT.LEFT,
	)


#============================================
def build_document(input_files: list[str], output_path: str, title: str, date: str) -> None:
	"""Build a DOCX from one or more cleaned HTML files."""
	if not output_path.endswith(".docx"):
		raise ValueError(f"Output file must have .docx extension, got: {output_path}")
	if os.path.exists(output_path):
		raise FileExistsError(f"Output file already exists: {output_path}")
	styles = ef_tools.style_loader.load_styles()
	doc = docx.Document()
	ef_tools.docx_builder.setup_styles(doc, styles)
	section = doc.sections[0]
	page_margin = styles["page"]["margin"]
	section.top_margin = docx.shared.Inches(page_margin)
	section.bottom_margin = docx.shared.Inches(page_margin)
	section.left_margin = docx.shared.Inches(page_margin)
	section.right_margin = docx.shared.Inches(page_margin)
	ef_tools.docx_builder.setup_header(doc, section, date, styles)
	total_questions = 0
	for input_file in input_files:
		with open(input_file, "r", encoding="utf-8") as handle:
			html_text = handle.read()
		html_doc = lxml.html.fromstring(html_text)
		total_questions += len(html_doc.xpath(TAKE_QUESTION_XPATH))
	add_front_page(doc, title, date, total_questions, styles)
	question_number = 1
	prev_element = "chapter"
	for input_file in input_files:
		document_title = extract_document_title(input_file)
		heading = doc.add_paragraph()
		heading.style = doc.styles["Exam Heading 2"]
		ef_tools.docx_builder.add_rich_text_runs(heading, document_title)
		with open(input_file, "r", encoding="utf-8") as handle:
			html_text = handle.read()
		html_doc = lxml.html.fromstring(html_text)
		for question_div in html_doc.xpath(TAKE_QUESTION_XPATH):
			prev_element = add_question_body(
				doc,
				input_file,
				question_div,
				question_number,
				styles,
				prev_element,
			)
			question_number += 1
	doc.save(output_path)


#============================================
def main() -> None:
	"""Main entry point."""
	args = parse_args()
	build_document(args.input_files, args.output_file, args.title, args.date)
	print(f"Exam DOCX written to {args.output_file}")


#============================================
if __name__ == "__main__":
	main()
