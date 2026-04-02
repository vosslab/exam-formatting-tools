"""Tests for the DOCX exam builder."""

# Standard Library
import os

# Pip Modules
import docx

# Local Repo Modules
import docx_exam_builder


#============================================
def test_decode_html_entities():
	"""HTML entities should be decoded to Unicode."""
	assert docx_exam_builder.decode_html_entities("&Delta;G") == "\u0394G"
	assert docx_exam_builder.decode_html_entities("&alpha;-helix") == "\u03b1-helix"
	assert docx_exam_builder.decode_html_entities("plain text") == "plain text"


#============================================
def test_strip_number_prefix():
	"""Number prefixes should be stripped from statements."""
	assert docx_exam_builder.strip_number_prefix("22) Which of") == "Which of"
	assert docx_exam_builder.strip_number_prefix("3. The answer") == "The answer"
	assert docx_exam_builder.strip_number_prefix("No prefix here") == "No prefix here"


#============================================
def test_auto_layout_for_choices():
	"""Auto-layout returns (tab_style, items_per_row) tuples."""
	# 2 short choices -> Choices4 tabs, 2 per row
	assert docx_exam_builder.auto_layout_for_choices(['a', 'b']) == (4, 2)
	# 3 short choices -> Choices3 tabs, 3 per row (all on one line)
	assert docx_exam_builder.auto_layout_for_choices(['a', 'b', 'c']) == (3, 3)
	# 4 short choices -> Choices4 tabs, 4 per row
	assert docx_exam_builder.auto_layout_for_choices(['a', 'b', 'c', 'd']) == (4, 4)
	# 4 long choices -> Choices4 tabs, 2 per row (2+2, avoids 3+1 orphan)
	long_4 = ['a very long choice text here!', 'another long one', 'third long', 'fourth long']
	assert docx_exam_builder.auto_layout_for_choices(long_4) == (4, 2)
	# 5 short choices -> Choices5 tabs, 5 per row
	assert docx_exam_builder.auto_layout_for_choices(['a', 'b', 'c', 'd', 'e']) == (5, 5)
	# 5 medium choices -> Choices3 tabs, 3 per row (3+2, avoids 4+1 orphan)
	med_5 = ['abcdefghijklmnopqrst'] * 5
	assert docx_exam_builder.auto_layout_for_choices(med_5) == (3, 3)


#============================================
def test_select_question_style():
	"""Question style should vary based on previous element."""
	assert docx_exam_builder.select_question_style('question') == 'Question Heading'
	assert docx_exam_builder.select_question_style('choices') == 'Question Heading'
	assert docx_exam_builder.select_question_style('chapter') == 'Question Follow'
	assert docx_exam_builder.select_question_style('image') == 'Question Follow'
	assert docx_exam_builder.select_question_style('table') == 'Question Follow'


#============================================
def test_overwrite_protection(tmp_path):
	"""Builder should refuse to overwrite existing files."""
	output_path = str(tmp_path / "existing.docx")
	# create the file first
	with open(output_path, 'w') as f:
		f.write("dummy")
	exam_data = {'title': 'Test', 'sections': []}
	raised = False
	# should raise FileExistsError
	try:
		docx_exam_builder.build_document(exam_data, output_path)
	except FileExistsError:
		raised = True
	assert raised is True


#============================================
def test_build_minimal_document(tmp_path):
	"""Build a minimal exam and verify basic structure."""
	output_path = str(tmp_path / "minimal.docx")
	exam_data = {
		'title': 'Test Exam',
		'date': '2026-04-02',
		'sections': [{
			'heading': 'Multiple Choice',
			'questions': [{
				'statement': 'What is 2+2?',
				'choices': ['3', '4', '5'],
			}],
		}],
	}
	docx_exam_builder.build_document(exam_data, output_path)
	assert os.path.isfile(output_path)
	# open and verify content
	doc = docx.Document(output_path)
	# should have paragraphs for: title, name, score, heading, question, choices
	assert len(doc.paragraphs) >= 5
	# check styles exist
	style_names = [s.name for s in doc.styles]
	assert 'Question Heading' in style_names
	assert 'Question Follow' in style_names
	assert 'Chapter Heading' in style_names
	assert 'Choices' in style_names


#============================================
def test_build_with_image(tmp_path):
	"""Build an exam with an image and verify it embeds."""
	# create a tiny PNG image
	import PIL.Image
	img_path = str(tmp_path / "test_img.png")
	img = PIL.Image.new('RGB', (100, 50), color='red')
	img.save(img_path)
	output_path = str(tmp_path / "with_image.docx")
	exam_data = {
		'title': 'Image Test',
		'sections': [{
			'questions': [{
				'statement': 'What is shown?',
				'image': img_path,
				'choices': ['A thing', 'Another thing'],
			}],
		}],
	}
	docx_exam_builder.build_document(exam_data, output_path)
	assert os.path.isfile(output_path)
	doc = docx.Document(output_path)
	# verify inline shapes (images) exist
	assert len(doc.inline_shapes) >= 1
