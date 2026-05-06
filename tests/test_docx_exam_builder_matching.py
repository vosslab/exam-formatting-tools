"""Smoke test for matching-question rendering in docx_exam_builder.

Exercises the prompts_list/choices_list path end-to-end by feeding a tiny
synthetic YAML structure through build_document and inspecting the output.
"""

import pytest
import yaml
import docx

import docx_exam_builder


#============================================
def _matching_yaml() -> dict:
	"""Synthetic exam_data with a single 3-prompt matching question."""
	exam_data = {
		'title': 'Test Exam',
		'date': '2026-05-06',
		'sections': [
			{
				'heading': 'Section A',
				'questions': [
					{
						'statement': 'Match each term with its definition.',
						'prompts_list': [
							'<b>alpha</b>',
							'<b>beta</b>',
							'<b>gamma</b>',
						],
						'choices_list': [
							'first definition',
							'second definition',
							'third definition',
						],
					},
				],
			},
		],
	}
	return exam_data


#============================================
def test_matching_question_renders_prompts_blanks_and_choice_letters(tmp_path):
	"""prompts_list drives Q1-3. plus three blanks; choices_list renders A/B/C."""
	exam_data = _matching_yaml()
	output_path = tmp_path / "matching.docx"
	docx_exam_builder.build_document(exam_data, str(output_path))
	doc = docx.Document(str(output_path))
	body_text = "\n".join(p.text for p in doc.paragraphs)
	# multi-prompt span header
	assert 'Q1-3.' in body_text
	# numbered blanks: each prompt consumes one slot
	assert '___ 1.' in body_text
	assert '___ 2.' in body_text
	assert '___ 3.' in body_text
	# all prompt texts present (rich text strips the <b> tags into runs)
	assert 'alpha' in body_text
	assert 'beta' in body_text
	assert 'gamma' in body_text
	# lettered choices appear with (A)/(B)/(C) prefixes
	assert '(A)' in body_text
	assert '(B)' in body_text
	assert '(C)' in body_text
	assert 'first definition' in body_text


#============================================
def test_matching_question_emits_no_choice_tables(tmp_path):
	"""Layout policy: matching choices_list must not become a docx table.

	Pairs the no-table check with a positive choice-rendered check so a
	regression that silently drops choices doesn't pass this test.
	"""
	exam_data = _matching_yaml()
	output_path = tmp_path / "matching_no_tables.docx"
	docx_exam_builder.build_document(exam_data, str(output_path))
	doc = docx.Document(str(output_path))
	body_text = "\n".join(p.text for p in doc.paragraphs)
	# the question has no `table` field -- nothing should produce a docx table
	assert doc.tables == []
	# choices must have actually rendered (guards against silent drops)
	assert '(A)' in body_text
	assert 'first definition' in body_text


#============================================
def test_matching_question_layout_override_respected(tmp_path):
	"""An explicit `layout` integer drives choices_list column count."""
	exam_data = _matching_yaml()
	exam_data['sections'][0]['questions'][0]['layout'] = 2
	output_path = tmp_path / "matching_layout.docx"
	docx_exam_builder.build_document(exam_data, str(output_path))
	doc = docx.Document(str(output_path))
	body_text = "\n".join(p.text for p in doc.paragraphs)
	# all three lettered options still render under the override
	assert '(A)' in body_text
	assert '(B)' in body_text
	assert '(C)' in body_text


#============================================
def test_legacy_matching_terms_key_raises(tmp_path):
	"""Legacy `matching_terms` schema must error loudly, not silently render."""
	exam_data = {
		'title': 'Legacy',
		'date': '2026-05-06',
		'sections': [{
			'heading': 'Section A',
			'questions': [{
				'statement': 'Match.',
				'matching_terms': ['a', 'b'],
			}],
		}],
	}
	output_path = tmp_path / "legacy.docx"
	with pytest.raises(ValueError, match="matching_terms"):
		docx_exam_builder.build_document(exam_data, str(output_path))


#============================================
def test_matching_question_advances_counter_by_prompt_span(tmp_path):
	"""A 3-prompt matching block followed by a question numbers the next as 4."""
	exam_data = _matching_yaml()
	exam_data['sections'][0]['questions'].append({
		'statement': 'A regular MC question.',
		'choices': ['one', 'two', 'three', 'four'],
	})
	output_path = tmp_path / "matching_counter.docx"
	docx_exam_builder.build_document(exam_data, str(output_path))
	doc = docx.Document(str(output_path))
	body_text = "\n".join(p.text for p in doc.paragraphs)
	assert 'Q1-3.' in body_text
	# next question after the 3-prompt span must be numbered 4
	assert '4. A regular MC question.' in body_text


#============================================
def test_yaml_safe_load_round_trip_matches_synthetic(tmp_path):
	"""Round-trip the synthetic exam_data through YAML to catch shape regressions."""
	exam_data = _matching_yaml()
	yaml_path = tmp_path / "exam.yaml"
	yaml_path.write_text(yaml.safe_dump(exam_data, sort_keys=False), encoding='utf-8')
	loaded = yaml.safe_load(yaml_path.read_text(encoding='utf-8'))
	question = loaded['sections'][0]['questions'][0]
	assert 'prompts_list' in question
	assert 'choices_list' in question
	assert 'matching_terms' not in question
