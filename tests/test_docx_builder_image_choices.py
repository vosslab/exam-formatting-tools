"""Test DOCX image choice tab-stop rendering (no tables)."""

import base64

import docx
import docx.shared

import ef_tools.docx_builder
import ef_tools.layout
import ef_tools.style_loader


#============================================
def _make_styled_doc():
	"""Create a docx Document with the exam styles pre-loaded."""
	styles = ef_tools.style_loader.load_styles()
	doc = docx.Document()
	ef_tools.docx_builder.setup_styles(doc, styles)
	return doc


# Smallest valid 1x1 transparent PNG, base64-encoded.
_PNG_BYTES = base64.b64decode(
	"iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNkYAAAAAYAAjCB0C8AAAAASUVORK5CYII="
)


#============================================
def _write_pngs(tmp_path, count):
	"""Write count tiny PNG files into tmp_path and return their paths."""
	paths = []
	for index in range(count):
		path = tmp_path / f"choice_{index}.png"
		path.write_bytes(_PNG_BYTES)
		paths.append(str(path))
	return paths


#============================================
def test_add_image_choices_tabbed_emits_no_tables(tmp_path):
	"""The function must not create any docx tables."""
	image_paths = _write_pngs(tmp_path, 4)
	choices = [{"text": "curve", "image": path} for path in image_paths]
	doc = _make_styled_doc()
	ef_tools.docx_builder.add_image_choices_tabbed(
		doc, choices, image_width=1.0,
	)
	assert len(doc.tables) == 0


#============================================
def test_add_image_choices_tabbed_uses_choices_n_style(tmp_path):
	"""Image-choice paragraphs must use a Choices N style; that style's tab
	stops drive column alignment. Adding paragraph-level tab stops on top
	of style-level stops causes Word to see two close-but-not-equal stops
	per column and image columns visibly drift -- so this test pins the
	style assignment and rejects redundant paragraph-level stops."""
	image_paths = _write_pngs(tmp_path, 4)
	choices = [{"text": "", "image": path} for path in image_paths]
	doc = _make_styled_doc()
	ef_tools.docx_builder.add_image_choices_tabbed(
		doc, choices, image_width=1.0,
	)
	# image paragraph is the last one (or second-to-last when alt text is
	# emitted); both should use Choices 4 with style-level tab stops
	for paragraph in doc.paragraphs[-3:]:
		if paragraph.style.name == "Choices 4":
			# style-level tab stops exist
			style_stops = list(paragraph.style.paragraph_format.tab_stops)
			assert len(style_stops) > 0
			# paragraph-level stops are NOT duplicated on top of style stops
			para_stops = list(paragraph.paragraph_format.tab_stops)
			assert len(para_stops) == 0


#============================================
def test_add_image_choices_tabbed_inlines_one_image_per_choice(tmp_path):
	"""Each choice with an image must produce one inline image in the doc."""
	image_paths = _write_pngs(tmp_path, 3)
	choices = [{"text": "", "image": path} for path in image_paths]
	doc = _make_styled_doc()
	ef_tools.docx_builder.add_image_choices_tabbed(
		doc, choices, image_width=1.0,
	)
	assert len(doc.inline_shapes) == 3


#============================================
def test_add_image_choices_tabbed_renders_letter_prefixes(tmp_path):
	"""The paragraph must contain bold (A) (B) (C) prefix runs in order."""
	image_paths = _write_pngs(tmp_path, 3)
	choices = [{"text": "", "image": path} for path in image_paths]
	doc = _make_styled_doc()
	ef_tools.docx_builder.add_image_choices_tabbed(
		doc, choices, image_width=1.0,
	)
	# letter prefixes live in the same combined paragraph as the images;
	# search across all paragraphs for the bold-prefix carrier
	prefix_texts = []
	for paragraph in doc.paragraphs:
		texts = [run.text for run in paragraph.runs if run.text.startswith("(")]
		if texts:
			prefix_texts = texts
			break
	# image-bearing choices use no trailing space after the letter so the
	# bold "(A)" hugs the image edge below it
	assert "(A)" in prefix_texts
	assert "(B)" in prefix_texts
	assert "(C)" in prefix_texts


#============================================
def test_add_image_choices_tabbed_mixed_text_and_images(tmp_path):
	"""Choices with no image must not produce an inline shape but still hold a column."""
	image_paths = _write_pngs(tmp_path, 2)
	choices = [
		{"text": "first", "image": image_paths[0]},
		{"text": "no picture", "image": None},
		{"text": "third", "image": image_paths[1]},
	]
	doc = _make_styled_doc()
	ef_tools.docx_builder.add_image_choices_tabbed(
		doc, choices, image_width=1.0,
	)
	# only the choices that supplied an image are embedded
	assert len(doc.inline_shapes) == 2
	# all three letter columns are still present in the caption paragraph;
	# its index depends on whether an alt-text paragraph was emitted, so
	# search across all paragraphs for the bold-prefix carrier
	prefix_texts = []
	for paragraph in doc.paragraphs:
		texts = [run.text for run in paragraph.runs if run.text.startswith("(")]
		if texts:
			prefix_texts = texts
			break
	# image-bearing choices render "(A)" without trailing space; the
	# text-only choice (B) keeps "(B) " with trailing space before its text
	assert {"(A)", "(B) ", "(C)"}.issubset(set(prefix_texts))


#============================================
def test_image_choice_max_width_per_cols_clamps_5col(tmp_path):
	"""5-col rows must be clamped to IMAGE_CHOICE_MAX_WIDTH_BY_COLS[5]
	even when the caller passes a much larger image_width. Without this
	cap, image E in a 5-col row pushes the cursor past its tab stop and
	the trailing image wraps to a new line."""
	image_paths = _write_pngs(tmp_path, 5)
	choices = [{"text": "", "image": path} for path in image_paths]
	doc = _make_styled_doc()
	# request 3.0" wide images; the per-col cap should clamp to 0.96"
	ef_tools.docx_builder.add_image_choices_tabbed(
		doc, choices, image_width=3.0, image_height=3.0,
	)
	cap_5col = ef_tools.docx_builder.IMAGE_CHOICE_MAX_WIDTH_BY_COLS[5]
	for shape in doc.inline_shapes:
		# rendered width must not exceed the per-col cap (with a small
		# tolerance for the height-bound aspect-preserve path)
		assert shape.width.inches <= cap_5col + 0.01


#============================================
def test_image_choice_max_width_per_cols_4col_differs_from_5col():
	"""The dict must give 4-col rows more horizontal budget than 5-col."""
	caps = ef_tools.docx_builder.IMAGE_CHOICE_MAX_WIDTH_BY_COLS
	assert caps[2] > caps[3] > caps[4] > caps[5]


#============================================
def test_zero_inline_image_margins_sets_dist_attrs_to_zero(tmp_path):
	"""Every inline image emitted by add_image_choices_tabbed must have
	distT/distB/distL/distR set to '0' on its <wp:inline> element. The
	default ~0.125" padding around each image steals visible width and
	breaks the tight column packing this layout relies on."""
	image_paths = _write_pngs(tmp_path, 3)
	choices = [{"text": "", "image": path} for path in image_paths]
	doc = _make_styled_doc()
	ef_tools.docx_builder.add_image_choices_tabbed(
		doc, choices, image_width=1.0,
	)
	WP = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
	inline_count = 0
	for shape in doc.inline_shapes:
		for inline in shape._inline.iter('{%s}inline' % WP):
			inline_count += 1
			for attr in ('distT', 'distB', 'distL', 'distR'):
				assert inline.get(attr) == '0', f"{attr}={inline.get(attr)!r}"
	assert inline_count >= 1


#============================================
def test_alt_text_paragraph_skipped_for_placeholder_image_text(tmp_path):
	"""When every choice carries the literal placeholder text 'image',
	the alt-text paragraph must be suppressed -- otherwise the row
	below the images becomes 'image image image image image' which is
	noise. This is the Q26 (histidine) shape in the real exam."""
	image_paths = _write_pngs(tmp_path, 5)
	choices = [{"text": "image", "image": path} for path in image_paths]
	doc = _make_styled_doc()
	before = len(doc.paragraphs)
	ef_tools.docx_builder.add_image_choices_tabbed(
		doc, choices, image_width=1.0,
	)
	# only the combined paragraph was added; no alt-text paragraph
	assert len(doc.paragraphs) - before == 1


#============================================
def test_alt_text_paragraph_emitted_for_meaningful_alt_text(tmp_path):
	"""When choices carry real captions ('right peak', 'down to up plot',
	etc.), the alt-text paragraph must be emitted below the image row.
	This is the Q80 (enzyme curves) shape in the real exam."""
	image_paths = _write_pngs(tmp_path, 5)
	captions = ["right peak", "down to up", "left peak", "up to down", "middle peak"]
	choices = [
		{"text": cap, "image": path}
		for cap, path in zip(captions, image_paths)
	]
	doc = _make_styled_doc()
	before = len(doc.paragraphs)
	ef_tools.docx_builder.add_image_choices_tabbed(
		doc, choices, image_width=1.0,
	)
	# combined paragraph + alt-text paragraph = 2 new paragraphs
	assert len(doc.paragraphs) - before == 2
	# the second new paragraph must contain the captions
	alt_para = doc.paragraphs[before + 1]
	alt_text = " ".join(run.text for run in alt_para.runs)
	for cap in captions:
		assert cap in alt_text


#============================================
def test_fit_picture_kwargs_picks_height_when_height_binds(tmp_path):
	"""When the source image is taller than wide and max_height is the
	tighter bound, fit_picture_kwargs must return height= (not width=)
	so aspect ratio is preserved without overflowing the height box."""
	# write a 10x40 PNG (4x taller than wide) so height-scale binds
	import PIL.Image
	tall_png = tmp_path / "tall.png"
	PIL.Image.new("RGB", (10, 40), color="white").save(tall_png)
	kwargs = ef_tools.docx_builder.fit_picture_kwargs(
		str(tall_png), max_width=1.0, max_height=1.0
	)
	assert "height" in kwargs and "width" not in kwargs


#============================================
def test_fit_picture_kwargs_picks_width_when_width_binds(tmp_path):
	"""Wide source image with width as the tighter bound: width= wins."""
	import PIL.Image
	wide_png = tmp_path / "wide.png"
	PIL.Image.new("RGB", (40, 10), color="white").save(wide_png)
	kwargs = ef_tools.docx_builder.fit_picture_kwargs(
		str(wide_png), max_width=1.0, max_height=1.0
	)
	assert "width" in kwargs and "height" not in kwargs


#============================================
def test_row_image_height_picks_smallest_binding_height(tmp_path):
	"""_row_image_height must return the minimum rendered height across
	all images so the row stays uniform when source aspect ratios differ.
	Construct two images whose rendered heights DIFFER so a buggy max()
	or `return max_height` cannot pass: the wide image (width-bound)
	produces 0.25, the tall image (height-bound) produces 1.0; the
	correct minimum is 0.25."""
	import PIL.Image
	# wide image (w:h = 4:1): width-bound -> h_at_max_w = 1.0 * 1/4 = 0.25
	wide_png = tmp_path / "wide.png"
	PIL.Image.new("RGB", (400, 100), color="white").save(wide_png)
	# tall image (w:h = 1:4): height-bound -> capped at max_height = 1.0
	tall_png = tmp_path / "tall.png"
	PIL.Image.new("RGB", (100, 400), color="white").save(tall_png)
	choices = [
		{"text": "", "image": str(wide_png)},
		{"text": "", "image": str(tall_png)},
	]
	# the wide image is the binding constraint; min(0.25, 1.0) = 0.25
	result = ef_tools.docx_builder._row_image_height(choices, 1.0, 1.0)
	assert abs(result - 0.25) < 0.001
