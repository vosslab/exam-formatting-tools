"""Lock the invariant that no rendered choice paragraph carries the
abstract base style name 'Choice'.

The 'Choice' style is registered in setup_styles as the parent of
'Choices 2', 'Choices 3', 'Choices 4', and 'Choices 5'. It must
remain abstract -- only the concrete children should ever land on a
paragraph. Three previous bugs collapsed paragraphs to bare 'Choice':
the CHOICES_STYLE_NAME[1] mapping, an override in
add_choices_paragraph triggered by any image-bearing choice, and a
hardcoded assignment in add_image_choices_tabbed.
"""

import os
import re
import sys
import base64

import yaml
import docx

import git_file_utils
import ef_tools.docx_builder
import ef_tools.style_loader

# put the repo root on sys.path so yaml_to_exam_docx.py is importable
_REPO_ROOT = git_file_utils.get_repo_root()
if _REPO_ROOT not in sys.path:
	sys.path.insert(0, _REPO_ROOT)
import yaml_to_exam_docx


_CHOICES_N_PATTERN = re.compile(r"^Choices [2-5]$")

# smallest valid 1x1 transparent PNG, base64-encoded
_PNG_BYTES = base64.b64decode(
	"iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNkYAAAAAYAAjCB0C8AAAAASUVORK5CYII="
)


#============================================
def _make_styled_doc():
	"""Create a docx Document with the exam styles pre-loaded."""
	styles = ef_tools.style_loader.load_styles()
	doc = docx.Document()
	ef_tools.docx_builder.setup_styles(doc, styles)
	return doc


#============================================
def _write_png(tmp_path, name):
	"""Write a tiny PNG to tmp_path/name and return its path."""
	path = tmp_path / name
	path.write_bytes(_PNG_BYTES)
	return str(path)


#============================================
def test_text_choices_use_concrete_choices_n_style():
	"""Four short text choices land on a Choices [2-5] style."""
	doc = _make_styled_doc()
	choices = ["yes", "no", "maybe", "unknown"]
	ef_tools.docx_builder.add_choices_paragraph(
		doc, choices, tab_style=4, items_per_row=4,
	)
	style_name = doc.paragraphs[-1].style.name
	assert _CHOICES_N_PATTERN.match(style_name), style_name


#============================================
def test_long_text_choices_avoid_bare_choice_style():
	"""Single-column (vertical stack) layout still uses Choices N, not Choice."""
	doc = _make_styled_doc()
	# tab_style=1 used to map to bare 'Choice'; it now maps to Choices 2
	choices = ["a long answer", "another long answer"]
	ef_tools.docx_builder.add_choices_paragraph(
		doc, choices, tab_style=1, items_per_row=1,
	)
	style_name = doc.paragraphs[-1].style.name
	assert _CHOICES_N_PATTERN.match(style_name), style_name


#============================================
def test_mixed_text_image_choices_avoid_bare_choice_style(tmp_path):
	"""Mixed text+image choices route through Choices N, not bare Choice."""
	image_path = _write_png(tmp_path, "mix.png")
	choices = [
		{"text": "first option"},
		{"text": "second option with image", "image": image_path},
		{"text": "third option"},
	]
	doc = _make_styled_doc()
	ef_tools.docx_builder.add_choices_paragraph(
		doc, choices, tab_style=3, items_per_row=3, image_width=1.0,
	)
	style_name = doc.paragraphs[-1].style.name
	assert _CHOICES_N_PATTERN.match(style_name), style_name


#============================================
def test_all_image_choices_avoid_bare_choice_style(tmp_path):
	"""add_image_choices_tabbed lands on a Choices [2-5] style."""
	images = [_write_png(tmp_path, f"img_{i}.png") for i in range(4)]
	choices = [{"text": "", "image": path} for path in images]
	doc = _make_styled_doc()
	ef_tools.docx_builder.add_image_choices_tabbed(
		doc, choices, image_width=1.0,
	)
	style_name = doc.paragraphs[-1].style.name
	assert _CHOICES_N_PATTERN.match(style_name), style_name


#============================================
def test_choices_n_styles_inherit_from_choice_base():
	"""Each Choices N style declares Choice as its base_style.

	This locks the inheritance edge that motivates keeping Choice as a
	registered (but never paragraph-applied) abstract style.
	"""
	doc = _make_styled_doc()
	choices_two = doc.styles["Choices 2"]
	assert choices_two.base_style is not None
	assert choices_two.base_style.name == "Choice"


#============================================
def test_matching_choices_list_uses_multi_column_style(tmp_path):
	"""A matching question's choices_list lands on a Choices [2-5] style.

	Goes through the full builder render path (`build_document`), which
	is the matching code path in yaml_to_exam_docx.py:202-208.
	"""
	# minimal YAML structure with one matching question; choices are short
	# chemistry formulae so the visible-width score allows a multi-column
	# layout (the regression collapsed this to bare Choice)
	exam_data = {
		"title": "Matching Render Path Test",
		"sections": [
			{
				"heading": "Test",
				"questions": [
					{
						"statement": "Match each bond type with an example.",
						"prompts_list": [
							"<b>non-polar covalent</b>",
							"<b>ionic</b>",
							"<b>hydrogen</b>",
							"<b>polar covalent</b>",
						],
						"choices_list": [
							"H<sub>2</sub>O",
							"CO<sub>2</sub>",
							"C &#8226; C",
							"C &#8801; C",
						],
					},
				],
			},
		],
	}
	output_path = str(tmp_path / "matching.docx")
	yaml_to_exam_docx.build_document(exam_data, output_path)
	doc = docx.Document(output_path)
	# find the choices_list paragraph: it carries the (A) prefix
	choices_paragraph = None
	for paragraph in doc.paragraphs:
		if paragraph.text.startswith("(A) "):
			choices_paragraph = paragraph
			break
	assert choices_paragraph is not None
	style_name = choices_paragraph.style.name
	assert _CHOICES_N_PATTERN.match(style_name), style_name


#============================================
def test_combined_yaml_no_paragraph_uses_choice_base_style(tmp_path):
	"""End-to-end: every paragraph in the combined exam DOCX is concrete.

	Builds Final_Exam/Final_Exam_2A_2B_combined.yaml from scratch, walks
	every paragraph in the result, and asserts none carry the abstract
	base style 'Choice'. This locks Gate G4 from the implementation plan.
	"""
	combined_yaml = os.path.join(
		_REPO_ROOT, "Final_Exam", "Final_Exam_2A_2B_combined.yaml")
	with open(combined_yaml, encoding="utf-8") as handle:
		exam_data = yaml.safe_load(handle)
	output_path = str(tmp_path / "combined.docx")
	yaml_to_exam_docx.build_document(exam_data, output_path)
	doc = docx.Document(output_path)
	bare_choice_paragraphs = [
		paragraph for paragraph in doc.paragraphs
		if paragraph.style and paragraph.style.name == "Choice"
	]
	assert bare_choice_paragraphs == []
