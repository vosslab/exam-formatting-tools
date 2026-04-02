#!/usr/bin/env python3
"""Generate a styled DOCX exam document from structured YAML input.

Uses python-docx to create a Word document with exam styles, auto-numbering,
bold choice prefixes, images, tables, and page headers. Reads the same
YAML format as odt_exam_builder.py (see docs/YAML_EXAM_FORMAT.md).
"""

# Standard Library
import os
import re
import html
import argparse

# Pip Modules
import yaml
import docx
import docx.shared
import docx.enum.text
import docx.enum.style
import docx.oxml.ns
import docx.oxml

# Local Repo Modules
import exam_defaults


#============================================
def parse_args() -> argparse.Namespace:
	"""Parse command-line arguments.

	Returns:
		Parsed argument namespace.
	"""
	parser = argparse.ArgumentParser(
		description="Generate a styled DOCX exam document from YAML input"
	)
	parser.add_argument(
		'-i', '--input', dest='input_file', required=True,
		help="Input YAML file with exam data"
	)
	parser.add_argument(
		'-o', '--output', dest='output_file', required=True,
		help="Output DOCX file path (must not already exist)"
	)
	args = parser.parse_args()
	return args


#============================================
def decode_html_entities(text: str) -> str:
	"""Decode HTML entities to Unicode characters for document output.

	YAML files use ASCII HTML entities (e.g., &Delta;, &alpha;) for
	special characters. This converts them to Unicode for rendering.

	Args:
		text: Text that may contain HTML entities.

	Returns:
		Text with entities decoded to Unicode.
	"""
	decoded = html.unescape(text)
	return decoded


#============================================
def strip_number_prefix(text: str) -> str:
	"""Strip leading question number prefix from statement text.

	Removes patterns like '22) ', '22. ', '3. ' from the start.

	Args:
		text: Statement text that may have a number prefix.

	Returns:
		Text with number prefix removed.
	"""
	stripped = re.sub(r'^\d+[).]\s*', '', text)
	return stripped


#============================================
def auto_layout_for_choices(choices: list) -> tuple:
	"""Determine the best layout for choices based on count and text length.

	Returns a (tab_style, items_per_row) tuple. The tab_style selects which
	set of tab stops to use (3, 4, or 5). The items_per_row controls how many
	choices go on each line. These can differ: e.g., tab_style=4 with
	items_per_row=2 gives a "Choices 2" layout using Choices 4 tab stops.

	Layout rules to avoid orphan rows (1 item on last row):
	- 2 items: Choices 4 tabs, 2 per row (even spacing via 4 % 2 == 0)
	- 3 items: Choices 3 (all on one row) if they fit; else vertical
	- 4 items: Choices 4 if they fit; else 2+2 (Choices 4 tabs, 2 per row);
	  never Choices 3 which gives 3+1 orphan
	- 5 items: Choices 5 if they fit; Choices 3 gives 3+2 (no orphan);
	  skip Choices 4 (4+1 orphan) and Choices 2 (2+2+1 orphan)

	Args:
		choices: List of choice text strings.

	Returns:
		Tuple of (tab_style, items_per_row).
	"""
	# empirically measured with bold "(A) " prefix included, minus 1 safety margin
	MAX_CHARS_5 = 17
	MAX_CHARS_4 = 23
	MAX_CHARS_3 = 30
	MAX_CHARS_2 = 49
	count = len(choices)
	if count < 1:
		return (5, 5)
	max_len = max(len(c) for c in choices)

	if count == 2:
		# 2 items: dedicated Choices 2 layout
		if max_len <= MAX_CHARS_2:
			return (2, 2)
		# very long: vertical stack
		return (1, 1)

	if count == 3:
		# prefer all 3 on one row
		if max_len <= MAX_CHARS_3:
			return (3, 3)
		# too long: vertical stack
		return (1, 1)

	if count == 4:
		# prefer 4 on one row
		if max_len <= MAX_CHARS_4:
			return (4, 4)
		# doesn't fit 4/row: use 2+2 layout (NOT Choices 3 which gives 3+1 orphan)
		if max_len <= MAX_CHARS_2:
			return (2, 2)
		# very long: vertical stack
		return (1, 1)

	# 5+ choices
	if max_len <= MAX_CHARS_5:
		# all 5 on one row
		return (5, 5)
	if max_len <= MAX_CHARS_3:
		# Choices 3 gives 3+2 (no orphan) -- skip Choices 4 which gives 4+1 orphan
		return (3, 3)
	if max_len <= MAX_CHARS_2:
		# long text: still use 3+2 with Choices 3 columns
		return (3, 3)
	# very long: vertical stack
	return (1, 1)


#============================================
def count_total_questions(sections: list) -> int:
	"""Count total number of questions across all sections.

	Args:
		sections: List of section dicts from exam data.

	Returns:
		Total count of questions.
	"""
	total = 0
	for section in sections:
		questions = section.get('questions', [])
		total += len(questions)
	return total


#============================================
def select_question_style(prev_element: str) -> str:
	"""Select appropriate question paragraph style based on previous element.

	Args:
		prev_element: Type of the previous element.

	Returns:
		Style name: 'Question Heading' or 'Question Follow'.
	"""
	if prev_element in ('question', 'choices'):
		style_name = 'Question Heading'
	else:
		# prev_element is 'chapter', 'table', or 'image'
		style_name = 'Question Follow'
	return style_name


#============================================
# tab stop positions for each choices layout (in inches)
CHOICES_TAB_STOPS = {
	1: [],
	2: [3.65],
	3: [2.33, 4.67],
	4: [1.75, 3.50, 5.25],
	5: [1.40, 2.80, 4.20, 5.60],
}

# style name for each tab layout
CHOICES_STYLE_NAME = {
	1: 'Choice',
	2: 'Choices 2',
	3: 'Choices 3',
	4: 'Choices 4',
	5: 'Choices 5',
}

# also keep 'Choices' in the old dict key for backward compat with explicit layout overrides
# (user can set layout: 3 in YAML which maps to tab_style=3)


#============================================
def _set_font_with_fallback(style, primary: str, fallback: str) -> None:
	"""Set font name on a style with a fallback font via XML.

	python-docx only supports a single font name. This sets the primary
	font and adds the fallback as hAnsi/cs font for cross-platform
	compatibility (Liberation Sans on Linux, Arial on Windows/Mac).

	Args:
		style: A python-docx paragraph or character style.
		primary: Primary font name (e.g., 'Liberation Sans').
		fallback: Fallback font name (e.g., 'Arial').
	"""
	style.font.name = primary
	# set the fallback font on the underlying XML for non-ascii/hAnsi
	rpr = style.element.get_or_add_rPr()
	rfonts_tag = docx.oxml.ns.qn('w:rFonts')
	rfonts = rpr.find(rfonts_tag)
	if rfonts is None:
		rfonts = docx.oxml.OxmlElement('w:rFonts')
		rpr.insert(0, rfonts)
	rfonts.set(docx.oxml.ns.qn('w:ascii'), primary)
	rfonts.set(docx.oxml.ns.qn('w:hAnsi'), fallback)
	rfonts.set(docx.oxml.ns.qn('w:cs'), fallback)


#============================================
def setup_styles(doc: docx.Document) -> None:
	"""Create all named exam styles in the document.

	Defines paragraph styles matching the exam style spec in
	docs/ODT_EXAM_STYLES.md, using python-docx style API.

	Args:
		doc: The Document to add styles to.
	"""
	# modify the built-in Normal style as our base
	normal = doc.styles['Normal']
	normal.font.size = docx.shared.Pt(11)
	normal.paragraph_format.space_after = docx.shared.Pt(1)
	normal.paragraph_format.space_before = docx.shared.Pt(0)
	normal.paragraph_format.line_spacing = 1.1
	_set_font_with_fallback(normal, 'Liberation Sans', 'Arial')

	# customize built-in Heading 1 to match ODT: 16pt bold, Liberation Sans
	# ODT: parent=Heading (14pt), 115% font size = ~16pt, bold, keep-with-next
	# set font via XML to include fallback: "Liberation Sans;Arial"
	h1 = doc.styles['Heading 1']
	h1.font.size = docx.shared.Pt(16)
	h1.font.bold = True
	h1.font.italic = False
	h1.font.color.rgb = docx.shared.RGBColor(0x00, 0x00, 0x00)
	h1.paragraph_format.space_before = docx.shared.Inches(0.17)
	h1.paragraph_format.space_after = docx.shared.Inches(0.08)
	h1.paragraph_format.keep_with_next = True
	# set font name with fallback via XML (python-docx only supports single name)
	_set_font_with_fallback(h1, 'Liberation Sans', 'Arial')

	# Question Heading: bold italic, hanging indent, keep-with-next
	# matching ODT: margin-top=0.15in, margin-bottom=0.02in
	qh = doc.styles.add_style('Question Heading', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
	qh.base_style = normal
	qh.font.bold = True
	qh.font.italic = True
	qh.font.size = docx.shared.Pt(11)
	qh.paragraph_format.left_indent = docx.shared.Inches(0.2)
	qh.paragraph_format.first_line_indent = docx.shared.Inches(-0.2)
	qh.paragraph_format.space_before = docx.shared.Inches(0.15)
	qh.paragraph_format.space_after = docx.shared.Inches(0.02)
	qh.paragraph_format.keep_with_next = True

	# Question Follow: same as Question Heading but no space above
	# used after images/tables/headings where extra gap would be redundant
	qf = doc.styles.add_style('Question Follow', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
	qf.base_style = normal
	qf.font.bold = True
	qf.font.italic = True
	qf.font.size = docx.shared.Pt(11)
	qf.paragraph_format.left_indent = docx.shared.Inches(0.2)
	qf.paragraph_format.first_line_indent = docx.shared.Inches(-0.2)
	qf.paragraph_format.space_before = docx.shared.Pt(0)
	qf.paragraph_format.space_after = docx.shared.Inches(0.02)
	qf.paragraph_format.keep_with_next = True

	# Chapter Heading: 14pt bold, dark purple, keep-with-next
	ch = doc.styles.add_style('Chapter Heading', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
	ch.base_style = normal
	ch.font.size = docx.shared.Pt(14)
	_set_font_with_fallback(ch, 'Liberation Sans', 'Arial')
	ch.font.bold = True
	ch.font.color.rgb = docx.shared.RGBColor(0x66, 0x00, 0xCC)
	ch.paragraph_format.space_before = docx.shared.Pt(6)
	ch.paragraph_format.space_after = docx.shared.Pt(3)
	ch.paragraph_format.keep_with_next = True

	# Heading 2: 14pt bold italic, for major section labels
	h2 = doc.styles.add_style('Exam Heading 2', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
	h2.base_style = normal
	h2.font.size = docx.shared.Pt(14)
	_set_font_with_fallback(h2, 'Liberation Sans', 'Arial')
	h2.font.bold = True
	h2.font.italic = True
	h2.paragraph_format.space_before = docx.shared.Pt(6)
	h2.paragraph_format.space_after = docx.shared.Pt(3)
	h2.paragraph_format.keep_with_next = True

	# Choice: base style for all choice layouts (10pt, indented, no tab stops)
	# also used directly for single-column vertical layout
	choice_base = doc.styles.add_style('Choice', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
	choice_base.base_style = normal
	choice_base.font.size = docx.shared.Pt(10)
	choice_base.paragraph_format.left_indent = docx.shared.Inches(0.15)
	choice_base.paragraph_format.first_line_indent = docx.shared.Inches(-0.05)
	choice_base.paragraph_format.space_before = docx.shared.Pt(0)
	choice_base.paragraph_format.space_after = docx.shared.Pt(1)

	# Choices 2 through 5: inherit from Choice, tab stops set per-paragraph
	for n in range(2, 6):
		style = doc.styles.add_style(f'Choices {n}', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
		style.base_style = choice_base


#============================================
def add_page_number_field(paragraph) -> None:
	"""Add 'Page X of Y' field codes to a paragraph using raw XML.

	python-docx does not have native page number field support,
	so we insert the Word XML field elements directly.

	Args:
		paragraph: The paragraph to add page numbers to.
	"""
	# "Page " text
	run_page = paragraph.add_run("Page ")
	run_page.font.size = docx.shared.Pt(9)
	# PAGE field
	run1 = paragraph.add_run()
	run1.font.size = docx.shared.Pt(9)
	fld_simple_page = docx.oxml.OxmlElement('w:fldSimple')
	fld_simple_page.set(docx.oxml.ns.qn('w:instr'), ' PAGE ')
	run1._element.addnext(fld_simple_page)
	# " of " text
	run_of = paragraph.add_run(" of ")
	run_of.font.size = docx.shared.Pt(9)
	# NUMPAGES field
	run2 = paragraph.add_run()
	run2.font.size = docx.shared.Pt(9)
	fld_simple_total = docx.oxml.OxmlElement('w:fldSimple')
	fld_simple_total.set(docx.oxml.ns.qn('w:instr'), ' NUMPAGES ')
	run2._element.addnext(fld_simple_total)


#============================================
def setup_header(section, date_str: str) -> None:
	"""Configure page header on body pages with page number, date, and name.

	Sets up different first page header (empty) and body page header with
	'Page X of Y | date | First Name:___' layout using tab stops.

	Args:
		section: The document section to configure.
		date_str: Date string to display in header center.
	"""
	# enable different first page header (empty on page 1)
	section.different_first_page_header_footer = True

	# body page header: "Page X of Y [tab] date [tab] First Name:___"
	header = section.header
	header.is_linked_to_previous = False
	header_para = header.paragraphs[0]
	# add tab stops: center at 3.65in, right at 7.3in
	header_para.paragraph_format.tab_stops.add_tab_stop(
		docx.shared.Inches(3.65), docx.enum.text.WD_TAB_ALIGNMENT.CENTER
	)
	header_para.paragraph_format.tab_stops.add_tab_stop(
		docx.shared.Inches(7.3), docx.enum.text.WD_TAB_ALIGNMENT.RIGHT
	)
	# add page number fields
	add_page_number_field(header_para)
	# tab to center, add date
	run_tab1 = header_para.add_run("\t")
	run_tab1.font.size = docx.shared.Pt(9)
	if date_str:
		run_date = header_para.add_run(date_str)
		run_date.font.size = docx.shared.Pt(9)
	# tab to right, add name line
	run_tab2 = header_para.add_run("\t")
	run_tab2.font.size = docx.shared.Pt(9)
	run_name = header_para.add_run("First Name:_____________________________")
	run_name.font.size = docx.shared.Pt(9)

	# first page header is empty (just clear it)
	first_header = section.first_page_header
	# clear any default content
	for p in first_header.paragraphs:
		p.text = ""


#============================================
def add_choices_paragraph(doc: docx.Document, choices: list,
	tab_style: int, items_per_row: int) -> None:
	"""Add a tab-separated choices paragraph with bold letter prefixes.

	Creates a paragraph with (A) (B) (C) format, using tab characters
	to align columns. The tab_style selects tab stop positions (3, 4, 5)
	and items_per_row controls how many choices per line.

	Args:
		doc: The Document to add the paragraph to.
		choices: List of choice text strings (no letter prefixes).
		tab_style: Tab stop layout (3, 4, or 5) from CHOICES_TAB_STOPS.
		items_per_row: Number of choices per line (may differ from tab_style).
	"""
	para = doc.add_paragraph()
	# select the appropriate style for this layout
	style_name = CHOICES_STYLE_NAME[tab_style]
	para.style = doc.styles[style_name]
	# set tab stops for the chosen layout
	tab_positions = CHOICES_TAB_STOPS[tab_style]
	for pos in tab_positions:
		para.paragraph_format.tab_stops.add_tab_stop(docx.shared.Inches(pos))
	# vertical stack: one choice per line, no tabs
	if items_per_row <= 1:
		for i, choice_text in enumerate(choices):
			if i > 0:
				run_br = para.add_run()
				run_br.add_break()
			letter = chr(ord('A') + i)
			bold_run = para.add_run(f"({letter}) ")
			bold_run.bold = True
			bold_run.font.size = docx.shared.Pt(10)
			text_run = para.add_run(decode_html_entities(choice_text))
			text_run.font.size = docx.shared.Pt(10)
		return
	# multi-column layout
	for i, choice_text in enumerate(choices):
		letter = chr(ord('A') + i)
		# line break before new rows (after first row)
		if i > 0 and i % items_per_row == 0:
			run_br = para.add_run()
			run_br.add_break()
		# tab before this choice (except first in each row)
		if i > 0 and i % items_per_row != 0:
			para.add_run("\t")
		# bold letter prefix
		bold_run = para.add_run(f"({letter}) ")
		bold_run.bold = True
		bold_run.font.size = docx.shared.Pt(10)
		# choice text (normal weight)
		text_run = para.add_run(decode_html_entities(choice_text))
		text_run.font.size = docx.shared.Pt(10)


#============================================
def add_table(doc: docx.Document, columns: list, rows: list) -> None:
	"""Add a table with a styled header row.

	The header row uses bold centered text with a light gray background.

	Args:
		doc: The Document to add the table to.
		columns: List of column header strings.
		rows: List of row data (each row is a list of cell strings).
	"""
	num_rows = len(rows) + 1
	num_cols = len(columns)
	table = doc.add_table(rows=num_rows, cols=num_cols)
	table.style = 'Table Grid'
	# header row
	for ci, col_text in enumerate(columns):
		cell = table.rows[0].cells[ci]
		cell.text = decode_html_entities(col_text)
		# bold and center the header text
		para = cell.paragraphs[0]
		for run in para.runs:
			run.bold = True
		para.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
		# light gray background via XML shading
		shading = docx.oxml.OxmlElement('w:shd')
		shading.set(docx.oxml.ns.qn('w:fill'), 'F2F2F2')
		shading.set(docx.oxml.ns.qn('w:val'), 'clear')
		cell._element.get_or_add_tcPr().append(shading)
	# data rows
	for ri, row_data in enumerate(rows):
		for ci, cell_text in enumerate(row_data):
			table.rows[ri + 1].cells[ci].text = decode_html_entities(cell_text)


#============================================
def build_document(exam_data: dict, output_path: str) -> None:
	"""Build a complete DOCX exam document from YAML data.

	Creates all styles, page layout, headers, and body content.
	Refuses to overwrite an existing file.

	Args:
		exam_data: Dict from YAML input with exam structure.
		output_path: Path for the output DOCX file.
	"""
	# enforce .docx extension
	if not output_path.endswith('.docx'):
		raise ValueError(f"Output file must have .docx extension, got: {output_path}")
	# safety check: never overwrite existing files
	if os.path.exists(output_path):
		raise FileExistsError(f"Output file already exists: {output_path}")

	doc = docx.Document()

	# setup styles
	setup_styles(doc)

	# page layout: 0.6in margins all around
	section = doc.sections[0]
	section.top_margin = docx.shared.Inches(0.6)
	section.bottom_margin = docx.shared.Inches(0.6)
	section.left_margin = docx.shared.Inches(0.6)
	section.right_margin = docx.shared.Inches(0.6)

	# setup page header (empty on first page, name/date on body pages)
	date_str = exam_data.get('date', '')
	setup_header(section, date_str)

	# --- first page boilerplate ---
	# exam title
	title = exam_data.get('title', 'Exam')
	title_para = doc.add_heading(decode_html_entities(title), level=1)
	title_para.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER

	# name line shifted right with tab
	name_line = exam_data.get('student_line', exam_defaults.DEFAULT_NAME_LINE)
	name_para = doc.add_paragraph()
	name_para.add_run("\t" + name_line)
	# add right tab stop
	name_para.paragraph_format.tab_stops.add_tab_stop(
		docx.shared.Inches(3.5), docx.enum.text.WD_TAB_ALIGNMENT.LEFT
	)

	# score line shifted right with tab
	sections = exam_data.get('sections', [])
	total_points = exam_data.get('total_points', None)
	if total_points is None:
		total_points = count_total_questions(sections)
	num_sections = exam_data.get('scoring_sections', exam_defaults.DEFAULT_SCORING_SECTIONS)
	score_line = exam_defaults.format_score_line(total_points, num_sections)
	score_para = doc.add_paragraph()
	score_para.add_run("\t" + score_line)
	score_para.paragraph_format.tab_stops.add_tab_stop(
		docx.shared.Inches(3.5), docx.enum.text.WD_TAB_ALIGNMENT.LEFT
	)

	# --- sections and questions ---
	question_counter = 1
	prev_element = 'chapter'
	for section_data in sections:
		# major section heading (Heading 2 level)
		heading = section_data.get('heading', '')
		if heading:
			para = doc.add_paragraph(decode_html_entities(heading))
			para.style = doc.styles['Exam Heading 2']
			prev_element = 'chapter'
		# chapter heading (purple, level 3)
		chapter = section_data.get('chapter', '')
		if chapter:
			para = doc.add_paragraph(decode_html_entities(chapter))
			para.style = doc.styles['Chapter Heading']
			prev_element = 'chapter'
		# questions
		questions = section_data.get('questions', [])
		for question in questions:
			# get explicit number if provided, otherwise use counter
			if 'number' in question:
				question_counter = question['number']
			question_number = question_counter
			# build question statement with number prefix
			statement = question.get('statement', '')
			if statement:
				# strip existing number prefix and add our own
				stripped = strip_number_prefix(statement)
				q_text = f"{question_number}. {decode_html_entities(stripped)}"
				# select style based on previous element
				style_name = select_question_style(prev_element)
				para = doc.add_paragraph(q_text)
				para.style = doc.styles[style_name]
				prev_element = 'question'
			# image (before choices, after question text)
			image_path = question.get('image', None)
			if image_path is not None and os.path.isfile(image_path):
				# add image with auto aspect ratio, max 5in wide
				doc.add_picture(image_path, width=docx.shared.Inches(5))
				prev_element = 'image'
			# table
			table_data = question.get('table', None)
			if table_data is not None:
				columns = table_data['columns']
				rows = table_data['rows']
				add_table(doc, columns, rows)
				prev_element = 'table'
			# choices
			choices = question.get('choices', None)
			if choices is not None and len(choices) > 0:
				layout_override = question.get('layout', None)
				if layout_override is not None:
					# explicit override: tab_style = items_per_row = layout value
					tab_style = layout_override
					items_per_row = layout_override
				else:
					# auto-determine layout
					tab_style, items_per_row = auto_layout_for_choices(choices)
				add_choices_paragraph(doc, choices, tab_style, items_per_row)
				prev_element = 'choices'
			# increment question counter
			question_counter += 1

	# save document
	doc.save(output_path)


#============================================
def main():
	"""Main entry point for DOCX exam builder."""
	args = parse_args()
	# load YAML
	with open(args.input_file, 'r') as f:
		exam_data = yaml.safe_load(f)
	# build document
	build_document(exam_data, args.output_file)
	print(f"Exam DOCX written to {args.output_file}")


#============================================
if __name__ == '__main__':
	main()
